"""
IFRS 9 (K-IFRS 1109) 금융자산 분류 마법사 v5
─────────────────────────────────────────────
실행 방법:
    pip install streamlit pdfplumber python-docx
    streamlit run app.py

v5 신규 기능:
    · 다중 파일 업로드 (accept_multiple_files=True)
      - PDF / DOCX 계약서 + 내부 기안문 동시 업로드
      - 파일별 역할(계약서·기안문·기타) 자동 감지
      - 모든 파일 텍스트를 통합하여 분석 베이스로 활용
    · 사업모형(BM) 자동 추론 (infer_bm)
      - 계약 조건(중도해지·만기보유 확약·이자수익 구조) 스캔
      - 경영 의도(안정수익·유동성관리·시세차익) 키워드 스캔
      - AC / FVOCI / FVPL 추론 + 신뢰도 + 근거 문구 출력
      - 사용자 확인 시 s_bm 자동 세팅
"""

from __future__ import annotations
import io
import re

import streamlit as st
try:
    import pandas as pd
    _PD_OK = True
except ImportError:
    pd = None  # type: ignore
    _PD_OK = False
import pandas as pd

try:
    import pandas as pd
    _PD_OK = True
except ImportError:
    _PD_OK = False
    pd = None

# ── 선택적 임포트 ─────────────────────────────────────────────────────────────
try:
    import pdfplumber
    _PDF_OK = True
except ImportError:
    _PDF_OK = False

try:
    from docx import Document as _DocxDocument
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False


# ══════════════════════════════════════════════════════════════════════════════
# 0. 상수 및 데이터
# ══════════════════════════════════════════════════════════════════════════════

NAVY = "#1E3A8A"
NAVY_LIGHT = "#EFF6FF"
NAVY_MID = "#DBEAFE"
GREEN = "#065F46"
GREEN_LIGHT = "#D1FAE5"
BLUE_MID = "#1D4ED8"
BLUE_LIGHT = "#EFF6FF"

# ── SPPI 불충족 사례 딕셔너리 ─────────────────────────────────────────────────
SPPI_CASES_DICT: dict = {
    "case_F": {
        "id": "case_F", "label": "상품F — 전환사채", "category": "지분연동",
        "instrument_desc": "확정수량의 발행자 지분상품으로 전환가능한 채권", "sppi_fail": True,
        "reason": "계약상 현금흐름이 기본대여계약과 일관되지 않는 수익 반영 — 발행자 지분가치에 연계됨",
        "standard_ref": ["§B4.1.14", "§B4.1.7A"],
        "judgment_criteria": "금융자산 주계약 시 내재파생 분리불가(§4.3.2). 지분가치 연동=§B4.1.7A 위반",
    },
    "case_G": {
        "id": "case_G", "label": "상품G — 역변동금리", "category": "TVM역방향",
        "instrument_desc": "역변동금리(시장이자율과 반비례) 대여금", "sppi_fail": True,
        "reason": "이자금액이 원금잔액의 화폐 시간가치 대가가 아님. 금리상승 시 이자 감소",
        "standard_ref": ["§B4.1.14", "§4.1.3⑵"],
        "judgment_criteria": "이자율 방향 확인. 역방향이면 TVM 대가 결여 → FVPL",
    },
    "case_H": {
        "id": "case_H", "label": "상품H — 이자이연+복리미발생", "category": "이자이연",
        "instrument_desc": "영구금융상품 — 지급여력 부족 시 이자이연 가능, 이연이자에 복리 미발생", "sppi_fail": True,
        "reason": "이자이연 가능 + 이연이자 복리 미발생 → 이자가 TVM의 진정한 대가가 아님",
        "standard_ref": ["§B4.1.14", "§4.1.3⑵"],
        "judgment_criteria": "이연이자에 복리가 붙으면 SPPI 가능. 영구적 특성 자체는 불충족 이유 아님",
    },
    "case_I": {
        "id": "case_I", "label": "상품I — 탄소가격지수(시장추적)", "category": "비대여변수연동",
        "instrument_desc": "매 보고기간 시장 탄소가격지수 변동 추적하여 이자율 조정 대여금", "sppi_fail": True,
        "reason": "기본대여위험·원가가 아닌 탄소가격지수(비대여 변수)에 따라 현금흐름 변동",
        "standard_ref": ["§B4.1.14 상품I", "§B4.1.8A"],
        "judgment_criteria": "cf. 상품EA(탄소배출 목표달성 고정bp 조정): 유의적 차이 없으면 SPPI 가능(§B4.1.10A)",
    },
    "case_road": {
        "id": "case_road", "label": "유료도로 통행량 연동", "category": "기초자산성과연동",
        "instrument_desc": "차량 통행수가 많을수록 현금흐름이 증가하는 금융자산", "sppi_fail": True,
        "reason": "계약상 현금흐름이 비금융자산(유료도로) 사용량 성과에 연동 — 기본대여계약 불일치",
        "standard_ref": ["§B4.1.16", "§B4.1.7A"],
        "judgment_criteria": "비소구 특성 자체는 불충족 아님. look-through 결과 기초자산 성과 연동 여부 판단",
    },
    "case_equity_idx": {
        "id": "case_equity_idx", "label": "주가지수 연동 이자·원금", "category": "지분연동",
        "instrument_desc": "이자·원금이 주가·주가지수 변동에 연동되는 채무상품", "sppi_fail": True,
        "reason": "기본대여계약 무관 위험(주식시장 위험)에 노출. 레버리지 동반 가능",
        "standard_ref": ["§B4.1.7A", "§B4.1.9"],
        "judgment_criteria": "지분위험 익스포저 = §B4.1.7A 핵심원칙 위반",
    },
    "case_profit": {
        "id": "case_profit", "label": "이익참가사채", "category": "채무자성과연동",
        "instrument_desc": "이자가 채무자의 순이익·수익의 일정 비율로 결정되는 채무상품", "sppi_fail": True,
        "reason": "채무자 사업성과에 연동된 수익 반영 → 기본대여계약 불일치",
        "standard_ref": ["§B4.1.7A", "§B4.1.8A"],
        "judgment_criteria": "순전히 신용위험 변동 보상인 구조라면 예외적 SPPI 가능",
    },
    "case_leverage": {
        "id": "case_leverage", "label": "레버리지 내재 상품", "category": "레버리지",
        "instrument_desc": "독립 옵션, 선도계약, 스왑 등 레버리지 내재 금융자산", "sppi_fail": True,
        "reason": "레버리지는 현금흐름 변동성을 이자의 경제적 특성을 초과하여 높임 → 이자 성격 상실",
        "standard_ref": ["§B4.1.9"],
        "judgment_criteria": "캡·플로어는 레버리지 없으면 SPPI 가능(§B4.1.11). 독립 옵션·선도·스왑은 항상 레버리지",
    },
    "case_tvm": {
        "id": "case_tvm", "label": "TVM 변형 — 이자율 기간 불일치", "category": "TVM변형",
        "instrument_desc": "이자기산기간 불일치: 1년이자율로 매월 재설정 / 5년만기에 5년이자율로 6개월 재설정", "sppi_fail": True,
        "reason": "TVM 요소 변형. 벤치마크 현금흐름과 합리적 시나리오에서 유의적 차이 가능",
        "standard_ref": ["§B4.1.9B", "§B4.1.9C", "§B4.1.9D"],
        "judgment_criteria": "규제이자율 TVM 대용 예외(§B4.1.9E). 미미한 영향이면 무시(§B4.1.18)",
    },
    "case_nonrecourse": {
        "id": "case_nonrecourse", "label": "지분담보 비소구 대여금", "category": "기초자산성과연동",
        "instrument_desc": "지분포트폴리오 담보 비소구 대여금 — 지분가격 하락 시 은행 손실", "sppi_fail": True,
        "reason": "채권자가 지분가격 하락 위험(풋옵션 동일 경제효과) 부담. 현금흐름이 지분성과에 연동",
        "standard_ref": ["§B4.1.16A", "§B4.1.17"],
        "judgment_criteria": "비소구 자체는 FVPL 이유 아님. look-through 결과 지분위험 노출 → SPPI 불충족",
    },
    "case_embedded_fa": {
        "id": "case_embedded_fa", "label": "금융자산 주계약 내재파생", "category": "내재파생(FA주계약)",
        "instrument_desc": "채무상품 주계약에 내재된 지분연계 이자·원금 지급계약", "sppi_fail": True,
        "reason": "금융자산 주계약 → §4.3.2 분리 금지. 복합계약 전체 SPPI 테스트 시 지분연계로 불충족",
        "standard_ref": ["§4.3.2", "§B4.3.5⑶"],
        "judgment_criteria": "금융부채 주계약과의 핵심 차이: 금융자산 주계약에는 내재파생 분리 금지",
    },
    "case_A_ok": {
        "id": "case_A_ok", "label": "상품A — 인플레이션 연계 (SPPI 충족 참고)", "category": "참고:SPPI충족",
        "instrument_desc": "발행통화 인플레이션지수 연계, 비레버리지, 원금 보장 채권", "sppi_fail": False,
        "reason": "SPPI 충족: 인플레이션 연계는 TVM을 현행 수준으로 재설정 — TVM 대가에 해당",
        "standard_ref": ["§B4.1.13 상품A", "§B4.1.7A"],
        "judgment_criteria": "채무자 성과·주가지수 추가 연계 시 불충족. 비레버리지 조건 필수 확인",
    },
}

# ── 기준서 핵심 문구 사전 ─────────────────────────────────────────────────────
STD_TEXTS: dict = {
    "§4.1.1": "금융자산은 최초 인식 시 (a) 금융자산의 관리를 위한 사업모형, (b) 금융자산의 계약상 현금흐름 특성에 근거하여 상각후원가·FVOCI·FVPL로 분류한다.",
    "§4.3.2": "복합계약의 주계약이 이 기준서의 적용범위에 해당하는 자산이라면, 내재파생상품을 별도로 회계처리하지 않는다. 복합계약 전체에 §4.1.1~4.1.5를 적용한다.",
    "§4.3.3": "주계약이 이 기준서의 적용범위에 해당하는 자산이 아닌 경우, 내재파생상품은 ① 경제적 특성·위험이 주계약과 밀접하게 관련되지 않고 ② 별도 계약이라면 파생상품 정의를 충족하며 ③ 공정가치 변동을 당기손익으로 인식하지 않는 경우에 분리한다.",
    "§B4.1.7A": "이자는 화폐의 시간가치, 신용위험, 기타 기본대여 위험·원가, 이윤에 대한 대가로만 구성되어야 한다. 기본대여계약과 일관되지 않는 그 밖의 위험이나 변동성에 노출시키는 계약상 현금흐름은 SPPI를 충족하지 않는다.",
    "§B4.1.9": "레버리지는 현금흐름의 변동성을 증가시켜 이자의 경제적 특성을 상실하게 만든다. 독립적인 옵션, 선도, 스왑은 레버리지를 내재한다.",
    "§B4.1.9C": "TVM 변형 평가 시, 수정되지 않은 금융상품(기준 상품)의 현금흐름과 비교하여 미할인 현금흐름 차이가 유의적인지를 합리적으로 가능한 시나리오 범위에서 평가한다.",
    "§B4.1.21": "계약상 연계 트랑슈: (a) 트랑슈 자체 계약조건이 SPPI 충족, (b) 기초 금융상품 집합이 SPPI 특성 충족, (c) 트랑슈의 신용위험 익스포저가 기초집합의 신용위험을 초과하지 않을 때 SPPI를 충족할 수 있다.",
    "§B4.1.2B": "사업모형은 단순한 주장이 아닌 사실에 근거하여 결정된다. 평가 증거에는 금융자산의 성과가 어떻게 평가·보고되는지, 위험이 어떻게 관리되는지, 경영진이 어떻게 보상받는지가 포함된다.",
    "§4.1.5": "기업은 최초 인식 시점에 인식·측정 불일치(회계불일치)를 제거하거나 유의적으로 감소시키기 위해 금융자산을 FVPL로 측정하도록 취소불가로 지정할 수 있다.",
    "§4.1.4": "지분상품 투자에 대해 단기매매 목적이 아닌 경우, 최초 인식 시점에 공정가치 변동을 OCI에 표시하는 취소불가 선택을 할 수 있다. 이 경우 어떠한 손익도 나중에 당기손익으로 재분류되지 않는다.",
    "§5.7.5": "지분상품 FVOCI 지정 시, 처분·제거 시에도 OCI에 누적된 손익을 당기손익으로 재분류하지 않는다. 다만, 자본 내의 이전은 허용된다.",
    "§5.5 (ECL)": "기대신용손실 모형: 최초 인식 후 신용위험이 유의적으로 증가한 경우 전체기간 ECL을, 그렇지 않은 경우 12개월 ECL을 손실충당금으로 인식한다.",
}

# ── AI 분석 키워드 규칙 ──────────────────────────────────────────────────────
_AI_RULES: list = [
    ("s_asset","hybrid",["전환권","전환사채","전환청구권","전환가격","상환전환우선주","RCPS","신종자본증권","조건부전환","CoCo","코코본드"],"high","전환권·전환사채·신종자본증권 조항 감지","§4.3.2"),
    ("s_asset","equity",["보통주","우선주","출자금","주주","의결권","신주발행","지분율"],"high","지분·주주 관련 조항 감지","§4.1.4"),
    ("s_asset","deriv",["금리스왑","통화스왑","선물환","풋옵션","콜옵션","파생상품계약","이자율스왑"],"high","독립 파생상품 계약 조항 감지","§4.1.4"),
    ("s_asset","debt",["원금","이자율","사채","채권","대여금","만기","원리금","회사채","국채","ABS","MBS","CLO"],"medium","채권·대여금 관련 조항 감지","§4.1.2"),
    ("s_host","fa_host",["전환사채","전환권부사채","신주인수권부사채","BW","CB","EB","교환사채","신종자본증권","영구채"],"high","채권 주계약 복합계약 조항 감지","§4.3.2"),
    ("s_host","other_host",["금융부채","리스","임차","운용리스","상품공급계약","서비스계약"],"medium","비금융자산·금융부채 주계약 조항 감지","§4.3.3"),
    ("s_sppi1","fail_equity",["전환권","전환가격","주가연동","주가지수연동","주식전환","KOSPI","코스피연동"],"high","주가·지분가치 연동 현금흐름 조항 감지","§B4.1.14"),
    ("s_sppi1","fail_profit",["이익참가","순이익연동","수익연동","이익배분","성과연동이자","매출연동"],"high","채무자 순이익·수익 연동 이자 조항 감지","§B4.1.7A"),
    ("s_sppi1","fail_inverse",["역변동금리","inverse floater","인버스 플로터","시장금리 반비례"],"high","역변동금리 조항 감지","§B4.1.14"),
    ("s_sppi1","fail_defer",["이자이연","이자지급 유예","이자 미지급","영구채","발행자 재량 이자"],"high","이자이연·영구채 조항 감지","§B4.1.14"),
    ("s_sppi1","fail_commodity",["탄소가격","원자재연동","금가격연동","유가연동","탄소배출권"],"high","원자재·탄소가격지수 연동 조항 감지","§B4.1.8A"),
    ("s_sppi1","none",["고정금리","변동금리","SOFR","EURIBOR","CD금리","기준금리","원리금","이자지급일"],"medium","단순 고정·변동금리 원리금 구조 감지","§B4.1.7A"),
    ("s_sppi2","tvm_ok",["3개월 CD금리","91일물","3M SOFR","1개월","분기별 재설정","고정이자율"],"medium","이자기산기간 일치 구조 감지","§B4.1.9B"),
    ("s_sppi3","clause_ok",["중도상환","조기상환","만기 전 상환"],"medium","중도상환 조항 감지","§B4.1.10~12"),
    ("s_sppi4","tranche_pass",["ABS","MBS","CLO","CDO","CBO","유동화","특수목적법인","SPC","트랑슈","자산유동화","선순위","후순위"],"medium","유동화·트랑슈 구조 조항 감지","§B4.1.20~26"),
]
_AI_PRIORITY = {
    "s_asset": ["hybrid","equity","deriv","debt"],
    "s_host":  ["fa_host","other_host"],
    "s_sppi1": ["fail_equity","fail_profit","fail_inverse","fail_defer","fail_commodity","fail_leverage","none"],
    "s_sppi2": ["tvm_ok"],
    "s_sppi3": ["clause_ok"],
    "s_sppi4": ["tranche_pass","tranche_no"],
}
_CONF_CFG = {"high":("높음","#D1FAE5","#065F46"), "medium":("중간","#FEF3C7","#92400E"), "low":("낮음","#FEE2E2","#991B1B")}


# ══════════════════════════════════════════════════════════════════════════════
# 1. AI 분석 엔진
# ══════════════════════════════════════════════════════════════════════════════

def extract_text(file) -> str:
    """단일 파일 텍스트 추출 (내부 유틸, 오류 시 [오류] 접두어 문자열 반환)"""
    name = file.name.lower()
    if name.endswith(".pdf"):
        if not _PDF_OK:
            return "[오류] pdfplumber 미설치: pip install pdfplumber"
        try:
            data = file.read()
            pages = []
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                for pg in pdf.pages:
                    t = pg.extract_text()
                    if t:
                        pages.append(t)
            return "\n".join(pages)
        except Exception as e:
            return f"[오류] PDF 추출 실패: {e}"
    if name.endswith(".docx"):
        if not _DOCX_OK:
            return "[오류] python-docx 미설치: pip install python-docx"
        try:
            data = file.read()
            doc = _DocxDocument(io.BytesIO(data))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            return f"[오류] DOCX 추출 실패: {e}"
    return "[오류] PDF 또는 DOCX 파일만 지원합니다."


# ── 파일 역할 감지 키워드 ────────────────────────────────────────────────────
_ROLE_KEYWORDS = {
    "계약서": ["계약서", "사채", "채권", "약정서", "인수계약", "발행조건", "사채권자",
               "갑", "을", "제1조", "제2조", "계약 당사자", "bond", "agreement"],
    "기안문": ["기안", "결재", "검토의견", "투자목적", "투자계획", "자산 운용",
               "기대수익", "가입 사유", "보고", "품의", "업무보고", "사업모형",
               "포트폴리오 전략", "내부 기준"],
}

def _detect_role(filename: str, text: str) -> str:
    """파일명·내용으로 역할 추정: '계약서' | '기안문' | '기타'"""
    fn = filename.lower()
    # 파일명 우선
    for role, kws in _ROLE_KEYWORDS.items():
        if any(kw in fn for kw in kws[:5]):
            return role
    # 내용 기반 카운트
    scores = {}
    text_lower = text.lower()
    for role, kws in _ROLE_KEYWORDS.items():
        scores[role] = sum(1 for kw in kws if kw in text_lower)
    best = max(scores, key=scores.get)
    return best if scores[best] >= 2 else "기타"


def extract_text_from_files(files: list) -> list:
    """
    여러 파일에서 텍스트를 추출하여 파일별 dict 리스트로 반환.
    반환 구조: [{"filename": str, "role": str, "text": str, "error": bool}, ...]
    """
    results = []
    for f in files:
        try:
            # file_uploader는 재사용 불가 — BytesIO로 래핑
            raw = f.read()
            name_lower = f.name.lower()
            text = ""
            error = False

            if name_lower.endswith(".pdf"):
                if not _PDF_OK:
                    text = "[오류] pdfplumber 미설치: pip install pdfplumber"
                    error = True
                else:
                    try:
                        pages = []
                        with pdfplumber.open(io.BytesIO(raw)) as pdf:
                            for pg in pdf.pages:
                                t = pg.extract_text()
                                if t:
                                    pages.append(t)
                        text = "\n".join(pages) if pages else "[경고] 텍스트를 추출할 수 없습니다 (스캔 이미지 PDF 가능성)."
                    except Exception as e:
                        text = f"[오류] PDF 추출 실패: {e}"
                        error = True

            elif name_lower.endswith(".docx"):
                if not _DOCX_OK:
                    text = "[오류] python-docx 미설치: pip install python-docx"
                    error = True
                else:
                    try:
                        doc = _DocxDocument(io.BytesIO(raw))
                        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                    except Exception as e:
                        text = f"[오류] DOCX 추출 실패: {e}"
                        error = True
            else:
                text = "[오류] PDF 또는 DOCX 파일만 지원합니다."
                error = True

            role = _detect_role(f.name, text) if not error else "기타"
            results.append({"filename": f.name, "role": role, "text": text, "error": error})

        except Exception as e:
            results.append({"filename": getattr(f, "name", "unknown"), "role": "기타",
                            "text": f"[오류] 파일 처리 중 예외 발생: {e}", "error": True})
    return results


# ══════════════════════════════════════════════════════════════════════════════
# 1-B. BM 자동 추론 로직 (사업모형 추론)
# ══════════════════════════════════════════════════════════════════════════════

# AC 신호 — 계약 조건 키워드 (계약서 내)
_BM_AC_CONTRACT = [
    "중도해지 불가", "중도해지 제한", "만기 보유", "만기보유 확약", "만기까지 보유",
    "중도상환 불가", "조기상환 제한", "만기 일시 상환", "원리금 지급 조건",
    "이자 수취", "쿠폰 수취", "원금 상환", "신용위험 관리",
    "hold to maturity", "held to collect",
    "원금 회수", "고금리 상품", "확정 이자", "약정 이자율", "원리금 수취 확정",
    "만기 도래", "원리금 지급 일정", "이자 지급 주기", "약정 만기일",
    "중도 해지 불가", "해지 제한", "조기 해지 불가",
    # v7 — 샘플 파일 문구
    "고금리 상품 투자", "만기 보유 원칙", "이자 수익 확정",
]
_BM_AC_INTENT = [
    "장기적 안정 수익", "안정적 이자수익", "이자수익 확보", "이자수익 목적",
    "장기 보유", "만기까지 운용", "금리 고정", "신용 관리 목적",
    "원리금 수취", "이자 수익 구조", "현금흐름 수취", "만기 보유 전략",
    "AC 모형", "상각후원가",
    "안정적인 자금운용", "이자 수익", "안정 수익", "확정 수익 목적",
    "이자수익 중심", "이자 수입 확보", "금리 수익", "자금 운용 계획",
    "고금리 운용", "안정적 운용", "원금 보전", "원금 회수 목적",
    # v7 — 샘플 기안문 2 문구 그대로
    "안정적인 자금운용", "고금리 상품 투자를 통한", "안정적 자금 운용",
    "이자 수익을 통한", "원금 및 이자",
]
_BM_FVOCI_CONTRACT = [
    "유동성 조건", "수시 매매 가능", "중도 매각 가능", "필요 시 처분",
    "듀레이션 관리", "금리 변동 대응", "포트폴리오 재구성",
    "알엘씨 (LCR)", "유동성 커버리지", "자산부채관리", "ALM",
    "hold to collect and sell",
    "매각 가능성", "조건부 매각", "필요시 매각", "시장 매각",
    "중도 처분 가능", "매도 가능 조건", "이자 및 매각 수익",
    # v7
    "매각 가능성 보유", "선택적 매각",
]
_BM_FVOCI_INTENT = [
    "유동성 관리", "유동성 확보", "유동성 목적", "시장 상황에 따른 매각",
    "유동성 포트폴리오", "필요 시 매각", "ALM 목적", "자산부채 매칭",
    "만기 매칭", "이자수익 유지", "매도 병행", "FVOCI 모형",
    "기타포괄손익", "듀레이션 조정",
    "전략적 투자", "우군 확보", "전략적 관계", "우호 지분", "관계 유지",
    "이자 및 매각 수익", "매각 수익 병행", "필요에 따라 매각",
    "유동성 확보 목적", "자금 조달 대비",
    # v7 — 샘플 기안문 2 문구 ("우군 확보", "지분 경쟁")
    "우군 확보", "지분 경쟁", "지분경쟁", "적극적인 우군",
    "우군 확보 목적", "전략적 우군", "매각 가능성 보유",
]
_BM_FVPL_CONTRACT = [
    "단기 보유", "단기 운용", "트레이딩 목적", "시장 가격 수익",
    "매매차익", "가격 변동 수익", "단기 자금 운용",
    "fair value through profit", "trading book",
    "공정가치 평가", "지분 매각", "주식 매각", "시세 변동",
    "매각 차익 실현", "단기 처분",
    # v7
    "투자이익 창출 목적", "시세 차익 실현",
]
_BM_FVPL_INTENT = [
    "단기 시세 차익", "매각 차익", "가격 차익", "단기 매매", "트레이딩",
    "공정가치 기준 운용", "단기 자금 운용", "시장성 자산 운용",
    "시세차익 목적", "매매 목적", "투기 목적", "단기 수익",
    "FVPL 모형", "당기손익",
    "투자이익 창출", "단기 차익", "시세 변동", "지분 매각",
    "매각 이익", "자본 이득", "공정가치 평가 이익", "단기 투자 회수",
    # v7 — 샘플 기안문 1 문구 ("투자이익 창출")
    "투자이익 창출", "시세 차익", "단기 운용",
    "투자 이익 창출", "이익 창출 목적",
]

# RCPS·지분상품 특별 감지 키워드
_EQUITY_SIGNAL_KEYWORDS = [
    "상환전환우선주", "RCPS", "전환우선주", "보통주", "우선주",
    "주식", "지분", "출자", "주주", "보통주 전환",
    "의결권", "배당", "신주인수권",
]
_RCPS_KEYWORDS = [
    "상환전환우선주", "RCPS", "전환우선주", "전환권", "상환권",
    "조건부 전환", "전환가격", "전환비율",
]
# v7 — RCPS+AC 복합 감지용 (계약서에 RCPS + 기안문에 AC 의도 동시 존재 시 노트)
_AC_INTENT_IN_CONTRACT = [
    "안정적인 자금운용", "안정적 자금운용", "이자 수익",
    "고금리 상품", "원금 회수", "만기 보유",
]


def infer_bm(files_info: list) -> dict:
    """
    파일별 텍스트를 계약서 / 기안문으로 구분하여 사업모형을 추론. (v7 고도화)

    v7:
    - 샘플 문서 문구("안정적인 자금운용", "우군 확보", "투자이익 창출") 직접 인식
    - RCPS+AC 복합 감지: 계약서 RCPS + 기안문 AC 의도 동시 → sppi_attention_note
    - 신뢰도 기준 완화 (짧은 기안문도 판단 가능)
    - suggested_bm: 신뢰도 무관 최유력값 항상 반환
    """
    ac_score, fvoci_score, fvpl_score = 0.0, 0.0, 0.0
    evidence_lines = []
    repeat_counter: dict = {}
    is_equity_signal = False
    equity_keywords_found: list = []
    has_rcps_in_contract = False
    has_ac_intent_in_any = False

    _NEG = ["불가", "제한", "확약", "보장", "의무", "반드시", "필수"]

    def _cw(sentence: str, base: float) -> float:
        sl = sentence.lower()
        return base + 0.5 if any(p in sl for p in _NEG) else base

    def _scan(text: str, fname: str, role: str,
              ac_kws, fvoci_kws, fvpl_kws, weight: float = 1.0):
        nonlocal ac_score, fvoci_score, fvpl_score
        sentences = [s.strip() for s in re.split(r'[\n。.!?\n]+', text) if len(s.strip()) > 3]
        for s in sentences:
            sl = s.lower()
            for kw in ac_kws:
                if kw.lower() in sl:
                    w = _cw(s, weight)
                    ac_score += w
                    repeat_counter[(fname, "AC")] = repeat_counter.get((fname, "AC"), 0) + 1
                    evidence_lines.append({"source": fname, "role": role,
                        "text": s[:100].strip(), "signal": "AC", "keyword": kw, "weight": round(w, 2)})
                    break
            for kw in fvoci_kws:
                if kw.lower() in sl:
                    w = _cw(s, weight)
                    fvoci_score += w
                    repeat_counter[(fname, "FVOCI")] = repeat_counter.get((fname, "FVOCI"), 0) + 1
                    evidence_lines.append({"source": fname, "role": role,
                        "text": s[:100].strip(), "signal": "FVOCI", "keyword": kw, "weight": round(w, 2)})
                    break
            for kw in fvpl_kws:
                if kw.lower() in sl:
                    w = _cw(s, weight)
                    fvpl_score += w
                    repeat_counter[(fname, "FVPL")] = repeat_counter.get((fname, "FVPL"), 0) + 1
                    evidence_lines.append({"source": fname, "role": role,
                        "text": s[:100].strip(), "signal": "FVPL", "keyword": kw, "weight": round(w, 2)})
                    break

    for fi in files_info:
        if fi["error"] or not fi["text"]:
            continue
        role, text, fname = fi["role"], fi["text"], fi["filename"]
        tl = text.lower()

        for ekw in _EQUITY_SIGNAL_KEYWORDS:
            if ekw.lower() in tl and ekw not in equity_keywords_found:
                equity_keywords_found.append(ekw)
                is_equity_signal = True

        if role == "계약서" and any(k.lower() in tl for k in _RCPS_KEYWORDS):
            has_rcps_in_contract = True
        if any(k.lower() in tl for k in _AC_INTENT_IN_CONTRACT):
            has_ac_intent_in_any = True

        if role == "계약서":
            _scan(text, fname, role, _BM_AC_CONTRACT, _BM_FVOCI_CONTRACT, _BM_FVPL_CONTRACT, 1.0)
        elif role == "기안문":
            _scan(text, fname, role, _BM_AC_INTENT, _BM_FVOCI_INTENT, _BM_FVPL_INTENT, 1.5)
        else:
            _scan(text, fname, role,
                  _BM_AC_CONTRACT + _BM_AC_INTENT,
                  _BM_FVOCI_CONTRACT + _BM_FVOCI_INTENT,
                  _BM_FVPL_CONTRACT + _BM_FVPL_INTENT, 0.7)

    for (fname, sig), cnt in repeat_counter.items():
        if cnt >= 3:
            if sig == "AC":    ac_score += 1.0
            elif sig == "FVOCI": fvoci_score += 1.0
            elif sig == "FVPL":  fvpl_score += 1.0

    sppi_attention_note = None
    if has_rcps_in_contract and has_ac_intent_in_any:
        sppi_attention_note = (
            "⚠️ 계약서에서 RCPS·전환권이, 문서에서 '안정적 자금운용' 등 AC 의도가 동시 감지되었습니다. "
            "계약상 원리금 요건(SPPI) 충족 여부를 집중 검토해야 합니다. "
            "주계약이 채무상품이면 복합계약 전체에 SPPI 테스트를 적용합니다(§4.3.2)."
        )

    total = ac_score + fvoci_score + fvpl_score
    scores = {"hold": ac_score, "both": fvoci_score, "trading": fvpl_score}
    suggested_bm = max(scores, key=scores.get) if total > 0 else "ambiguous"

    if total == 0:
        return {
            "proposed_bm": "ambiguous", "suggested_bm": "ambiguous", "confidence": "low",
            "ac_score": 0.0, "fvoci_score": 0.0, "fvpl_score": 0.0,
            "evidence_lines": [],
            "summary": "사업모형 관련 키워드가 감지되지 않았습니다. 직접 선택해 주세요.",
            "is_equity_signal": is_equity_signal, "equity_keywords": equity_keywords_found,
            "sppi_attention_note": sppi_attention_note,
        }

    winner = suggested_bm
    max_s = scores[winner]
    second_s = sorted(scores.values())[-2]
    gap_ratio = (max_s - second_s) / max(total, 1)

    # v7: 기준 완화 — 짧은 실무 문서 대응
    if gap_ratio >= 0.15 and max_s >= 1.5:
        confidence = "high"
    elif gap_ratio >= 0.05 or max_s >= 1.0:
        confidence = "medium"
    else:
        confidence = "low"

    sig_map = {"hold": "AC", "both": "FVOCI", "trading": "FVPL"}
    sig = sig_map.get(winner, "AC")
    top_ev = sorted([e for e in evidence_lines if e["signal"] == sig],
                    key=lambda x: x["weight"], reverse=True)[:4]
    other_ev = [e for e in evidence_lines if e["signal"] != sig][:2]

    bm_ko = {"hold": "AC — 계약상 현금흐름 수취 모형",
              "both": "FVOCI — 수취+매도 병행 모형",
              "trading": "FVPL — 공정가치 실현 모형"}
    parts = []
    if ac_score > 0:    parts.append(f"AC {ac_score:.1f}점")
    if fvoci_score > 0: parts.append(f"FVOCI {fvoci_score:.1f}점")
    if fvpl_score > 0:  parts.append(f"FVPL {fvpl_score:.1f}점")
    summary = f"{bm_ko.get(winner,'?')} 추론 ({', '.join(parts)})"
    if is_equity_signal:
        summary += f" | ⚠️ 지분 키워드: {', '.join(equity_keywords_found[:3])}"
    if sppi_attention_note:
        summary += " | 🔍 SPPI 집중검토 필요"

    return {
        "proposed_bm": winner, "suggested_bm": suggested_bm, "confidence": confidence,
        "ac_score": ac_score, "fvoci_score": fvoci_score, "fvpl_score": fvpl_score,
        "evidence_lines": top_ev + other_ev, "summary": summary,
        "is_equity_signal": is_equity_signal, "equity_keywords": equity_keywords_found,
        "sppi_attention_note": sppi_attention_note,
    }




def ai_analyze(text: str, files_info: list | None = None) -> dict:
    """
    키워드 기반 SPPI + 자산 성격 분석 (기존).
    files_info 가 전달되면 BM 추론을 추가로 수행하여 결과에 포함.
    """
    hits: dict = {}
    for step_id, val, kws, conf, basis, std_ref in _AI_RULES:
        matched = [kw for kw in kws if kw.lower() in text.lower()]
        if matched:
            hits.setdefault(step_id, {}).setdefault(val, []).extend(matched)

    proposed, evidence, conflicts = {}, [], []
    for step_id, priority in _AI_PRIORITY.items():
        step_hits = hits.get(step_id, {})
        if not step_hits:
            continue
        matched_vals = [v for v in step_hits if step_hits[v]]
        if len(matched_vals) > 1:
            winner = next((v for v in priority if v in matched_vals), matched_vals[0])
            conflicts.append({"step_id": step_id, "message": f"복수 신호 감지: {', '.join(matched_vals)} → {winner} 선택 (수동 확인 권장)"})
        else:
            winner = matched_vals[0]
        cnt = len(step_hits.get(winner, []))
        conf_level = "high" if cnt >= 3 else ("medium" if cnt >= 1 else "low")
        rule = next((r for r in _AI_RULES if r[0] == step_id and r[1] == winner), (step_id, winner, [], "low", "자동 감지", "—"))
        proposed[step_id] = winner
        evidence.append({"step_id": step_id, "value": winner, "keywords_found": step_hits.get(winner, []), "basis": rule[4], "std_ref": rule[5], "confidence": conf_level})

    # BM 추론 — files_info 가 있을 때만 실행
    bm_inference = infer_bm(files_info) if files_info else None

    return {
        "proposed_answers": proposed,
        "evidence_items": evidence,
        "conflict_flags": conflicts,
        "skippable_steps": [s for s in proposed if s != "s_bm"],
        "bm_inference": bm_inference,   # ← v5 신규
    }


# ══════════════════════════════════════════════════════════════════════════════
# 2. 분류 결과 계산 로직
# ══════════════════════════════════════════════════════════════════════════════

def sppi_fail_result(fail_key: str, case_key) -> dict:
    msgs = {
        "fail_equity":    "주가·주가지수 연동 이자·원금. 기본대여계약과 일관되지 않는 위험에 노출됩니다.",
        "fail_commodity": "원자재·탄소가격지수(시장추적) 연동. 비대여 변수에 따라 현금흐름이 변동합니다.",
        "fail_profit":    "채무자 순이익·수익 비율 연동. 사업성과에 연동된 수익이 반영됩니다.",
        "fail_inverse":   "역변동금리 — 이자가 TVM 대가가 아닙니다. 금리상승 시 이자 감소.",
        "fail_leverage":  "레버리지 내재 — 현금흐름 변동성이 이자의 경제적 특성을 초과합니다.",
        "fail_defer":     "이자이연 가능 + 이연이자 복리 미발생. 단, 복리가 붙으면 SPPI 가능.",
        "tvm_fail":       "TVM 변형이 유의적. 벤치마크 현금흐름과 합리적 시나리오에서 유의적 차이.",
        "clause_fail":    "계약조건으로 원리금 불일치 현금흐름 발생 (§B4.1.12 예외 미해당).",
        "tranche_fail":   "트랑슈 Look-through 조건 불충족 또는 최초 인식시점에 평가 불가.",
    }
    refs_map = {
        "fail_equity":   ["§B4.1.7A","§B4.1.14","§4.1.4"], "fail_commodity":["§B4.1.8A","§B4.1.14 상품I","§4.1.4"],
        "fail_profit":   ["§B4.1.7A","§B4.1.8A","§4.1.4"], "fail_inverse":  ["§B4.1.14","§4.1.3⑵","§4.1.4"],
        "fail_leverage": ["§B4.1.9","§4.1.4"],              "fail_defer":    ["§B4.1.14","§4.1.3⑵","§4.1.4"],
        "tvm_fail":      ["§B4.1.9C","§B4.1.9D","§4.1.4"], "clause_fail":   ["§B4.1.10","§B4.1.16","§4.1.4"],
        "tranche_fail":  ["§B4.1.26","§B4.1.21","§4.1.4"],
    }
    return {
        "classification":"FVPL","label":"당기손익-공정가치 (FVPL) — SPPI 불충족","color":"red",
        "reason":"SPPI 테스트 불충족. "+msgs.get(fail_key,""),
        "refs":refs_map.get(fail_key,["§4.1.4"]),
        "ecl":False,"recycling":False,"recycling_note":None,
        "accounting":["공정가치로 최초 인식 및 후속 측정","모든 공정가치 변동 → 당기손익","ECL(손상) 적용 없음"],
        "warning":None,"case_key":case_key,
    }


def compute_result(ans: dict) -> dict:
    at = ans.get("s_asset")
    if at == "deriv":
        return {"classification":"FVPL","label":"당기손익-공정가치 (FVPL) — 독립 파생상품","color":"red",
                "reason":"독립 파생상품은 항상 FVPL로 측정합니다. 레버리지 내재로 SPPI 불충족이며, AC·FVOCI 분류 불가.",
                "refs":["§4.1.4","§B4.1.9"],"ecl":False,"recycling":False,"recycling_note":None,
                "accounting":["공정가치로 최초 인식 및 후속 측정","모든 공정가치 변동 → 당기손익","ECL(손상) 적용 없음"],"warning":None,"case_key":None}
    if at == "hybrid":
        host = ans.get("s_host")
        if host == "other_host":
            if ans.get("s_sep") == "sep_ok":
                return {"classification":"FVPL","label":"내재파생상품 FVPL 분리 + 주계약 별도 처리","color":"red",
                        "reason":"분리 3요건 충족. 내재파생 → FVPL, 주계약 → 관련 기준서 별도 처리.",
                        "refs":["§4.3.3","§4.3.4","§B4.3.5"],"ecl":False,"recycling":False,"recycling_note":None,
                        "accounting":["내재파생상품 → FVPL","주계약 → 관련 기준서 별도 처리","재평가 금지 [§B4.3.11]"],
                        "warning":"밀접관련 내재파생(§B4.3.8)은 분리하지 않습니다: 레버리지 없는 금리캡·플로어 등","case_key":"case_embedded_fa"}
            return {"classification":"FVPL","label":"복합계약 전체 → FVPL","color":"red",
                    "reason":"분리 3요건 미충족 또는 내재파생상품 측정 불가. 복합계약 전체를 FVPL로 측정합니다.",
                    "refs":["§4.3.6","§4.3.7"],"ecl":False,"recycling":False,"recycling_note":None,
                    "accounting":["복합계약 전체 → FVPL","공정가치 변동 전액 → 당기손익"],"warning":None,"case_key":"case_embedded_fa"}
    if at == "equity":
        if ans.get("s_eq_trade") == "trade_yes":
            return {"classification":"FVPL","label":"당기손익-공정가치 (FVPL) — 단기매매 지분","color":"red",
                    "reason":"단기매매 목적 지분상품은 FVPL. FVOCI 취소불가 선택권 행사 불가.",
                    "refs":["§4.1.4"],"ecl":False,"recycling":False,"recycling_note":None,
                    "accounting":["공정가치로 측정","모든 공정가치 변동 → 당기손익","배당 → 당기손익","ECL 적용 없음"],"warning":None,"case_key":None}
        if ans.get("s_eq_fvoci") == "fvoci_yes":
            return {"classification":"FVOCI","label":"기타포괄손익-공정가치 (FVOCI) — 지분 취소불가 지정","color":"blue",
                    "reason":"단기매매가 아닌 지분상품에 FVOCI 취소불가 선택권 행사.",
                    "refs":["§4.1.4","§5.7.5","§5.7.6","§B5.7.1"],"ecl":False,"recycling":False,
                    "recycling_note":"지분상품 FVOCI: 처분 시에도 OCI 누적손익이 P&L로 재분류(Recycling)되지 않습니다.",
                    "accounting":["공정가치 변동 전액 → OCI","배당(투자원가 회수 제외) → 당기손익 [§B5.7.1]",
                                  "처분 시 OCI → P&L 재분류 금지 (Recycling 없음) [§5.7.5]","ECL 규정 미적용"],
                    "warning":"처분 시 OCI → P&L 재분류 없음(Recycling 금지). ECL 인식하지 않음.","case_key":None}
        return {"classification":"FVPL","label":"당기손익-공정가치 (FVPL) — 지분 기본값","color":"orange",
                "reason":"FVOCI 선택권 미행사 비단기매매 지분상품은 FVPL로 측정.",
                "refs":["§4.1.4"],"ecl":False,"recycling":False,"recycling_note":None,
                "accounting":["공정가치로 측정","모든 공정가치 변동 → 당기손익","배당 → 당기손익"],"warning":None,"case_key":None}
    case_map = {"fail_equity":"case_equity_idx","fail_commodity":"case_I","fail_profit":"case_profit",
                "fail_inverse":"case_G","fail_leverage":"case_leverage","fail_defer":"case_H"}
    sp1 = ans.get("s_sppi1")
    if sp1 and sp1 != "none":
        return sppi_fail_result(sp1, case_map.get(sp1))
    if ans.get("s_sppi2") == "tvm_fail":  return sppi_fail_result("tvm_fail","case_tvm")
    if ans.get("s_sppi3") == "clause_fail": return sppi_fail_result("clause_fail","case_road")
    if ans.get("s_sppi4") == "tranche_fail": return sppi_fail_result("tranche_fail",None)
    bm = ans.get("s_bm")
    if bm == "ambiguous":
        return {"classification":"FVPL","label":"추가 검토 필요 → 잠정 FVPL","color":"orange",
                "reason":"입력 정보만으로 사업모형을 명확히 판단하기 어렵습니다. 내부보고체계·과거 매도이력·보상방식 등 추가 증거를 검토하십시오.",
                "refs":["§B4.1.2B","§B4.1.5"],"ecl":False,"recycling":False,"recycling_note":None,
                "accounting":["잠정 FVPL 처리 후 추가 검토","사업모형 증거 문서화 필요"],
                "warning":"확정 분류 전 주요경영진 결정 근거를 반드시 문서화하십시오.","case_key":None}
    if bm == "trading":
        return {"classification":"FVPL","label":"당기손익-공정가치 (FVPL) — 잔여범주","color":"red",
                "reason":"공정가치 기준 관리·평가 또는 단기매매 포트폴리오 → FVPL.",
                "refs":["§B4.1.5","§B4.1.6","§4.1.4"],"ecl":False,"recycling":False,"recycling_note":None,
                "accounting":["공정가치로 측정","모든 공정가치 변동 → 당기손익","ECL 적용 없음"],"warning":None,"case_key":None}
    if ans.get("s_fvo") == "fvo_yes":
        return {"classification":"FVPL","label":"당기손익-공정가치 (FVPL) — FVO 지정","color":"orange",
                "reason":"회계불일치 해소를 위한 FVO 지정. 최초 인식시점 취소불가.",
                "refs":["§4.1.5","§B4.1.29~32"],"ecl":False,"recycling":False,"recycling_note":None,
                "accounting":["공정가치로 측정","모든 공정가치 변동 → 당기손익","최초 인식시점 지정·취소불가","ECL 적용 없음"],
                "warning":"FVO 지정은 취소불가입니다. 회계불일치 해소 요건을 면밀히 검토하십시오.","case_key":None}
    if bm == "hold":
        return {"classification":"AC","label":"상각후원가 (AC)","color":"green",
                "reason":"SPPI 충족 + 계약상 현금흐름 수취 목적 사업모형 → AC 측정.",
                "refs":["§4.1.2","§4.1.2⑴","§B4.1.2C","§B4.1.3A"],"ecl":True,"recycling":False,"recycling_note":None,
                "accounting":["최초 인식: 공정가치 + 거래원가","후속 측정: 유효이자율법 상각후원가",
                               "이자수익 → 당기손익 (유효이자율법)","ECL(기대신용손실) 손상 모형 적용 [§5.5]",
                               "처분 시 장부금액과 수취대가 차이 → 당기손익"],"warning":None,"case_key":None}
    if bm == "both":
        return {"classification":"FVOCI","label":"기타포괄손익-공정가치 (FVOCI) — 채무상품","color":"blue",
                "reason":"SPPI 충족 + 수취와 매도 둘 다가 목적인 사업모형 → FVOCI.",
                "refs":["§4.1.2A","§4.1.2A⑴","§B4.1.4A","§B4.1.4C"],"ecl":True,"recycling":True,"recycling_note":None,
                "accounting":["최초 인식·후속 측정: 공정가치","이자수익·ECL·외환손익 → 당기손익",
                               "그 외 공정가치 변동 → OCI","처분 시 OCI 누적손익 → P&L 재분류 (Recycling) [§5.7.2]"],
                "warning":"채무상품 FVOCI는 처분 시 OCI 손익이 P&L로 재분류됩니다. 지분상품 FVOCI와의 핵심 차이입니다.","case_key":None}
    return {"classification":"FVPL","label":"당기손익-공정가치 (FVPL)","color":"red",
            "reason":"잔여범주 FVPL 적용.","refs":["§4.1.4"],"ecl":False,"recycling":False,"recycling_note":None,
            "accounting":["공정가치로 측정","모든 공정가치 변동 → 당기손익"],"warning":None,"case_key":None}


# ══════════════════════════════════════════════════════════════════════════════
# 3. 스텝 시퀀스 계산
# ══════════════════════════════════════════════════════════════════════════════

def get_step_sequence(ans: dict) -> list:
    at = ans.get("s_asset")
    if not at: return ["s_asset"]
    if at == "deriv": return ["s_asset"]
    if at == "equity":
        base = ["s_asset","s_eq_trade"]
        if ans.get("s_eq_trade") == "trade_no": base.append("s_eq_fvoci")
        return base
    if at == "hybrid":
        base = ["s_asset","s_host"]
        host = ans.get("s_host")
        if not host: return base
        if host == "fa_host": return base + _debt_sppi_seq(ans)
        return base + ["s_sep"]
    return ["s_asset"] + _debt_sppi_seq(ans)


def _debt_sppi_seq(ans: dict) -> list:
    seq = ["s_sppi1"]
    if not ans.get("s_sppi1") or ans.get("s_sppi1") != "none": return seq
    seq.append("s_sppi2")
    if ans.get("s_sppi2") == "tvm_fail": return seq
    seq.append("s_sppi3")
    if ans.get("s_sppi3") == "clause_fail": return seq
    seq.append("s_sppi4")
    if ans.get("s_sppi4") == "tranche_fail": return seq
    seq.append("s_bm")
    bm = ans.get("s_bm")
    if not bm or bm in ("trading","ambiguous"): return seq
    seq.append("s_fvo")
    return seq


def is_terminal(ans: dict) -> bool:
    seq = get_step_sequence(ans)
    return seq[-1] in ans


# ══════════════════════════════════════════════════════════════════════════════
# 4. STEP_DEFS (기존 유지 + ref_key 추가)
# ══════════════════════════════════════════════════════════════════════════════

STEP_DEFS: dict = {
    "s_asset": {
        "tag":"","title":"STEP 0 — 금융자산의 기본 성격을 알려주세요","ref":"§4.1.1","ref_key":"§4.1.1",
        "desc":"🏷️ **이 단계는 금융상품의 '혈액형'을 확인하는 단계입니다.** 같은 채권이라도 전환사채처럼 특별한 조건이 붙어있으면 다른 경로로 분류됩니다.\n\n보유하고 있는 금융자산이 아래 중 어떤 성격에 가장 가까운지 선택하세요.",
        "helper":None,
        "options":[
            ("debt","📄 채무상품 — 원금+이자 구조","만약 '원금과 이자를 돌려받는 계약'(대출채권·회사채·국채 등)이라면 이 항목을 선택하세요."),
            ("equity","📈 지분상품 — 주식·출자금","만약 '주주로서 회사의 일부를 소유하는 계약'(보통주·우선주·출자증권)이라면 이 항목을 선택하세요."),
            ("deriv","🔄 독립 파생상품 — 옵션·선도·스왑","만약 주계약 없이 단독으로 거래되는 옵션·금리스왑·선물환 등이라면 이 항목을 선택하세요."),
            ("hybrid","⚠️ 복합계약 — 채권 안에 파생이 숨어있는 구조","만약 전환사채·신종자본증권처럼 채권(주계약)과 파생(내재파생)이 결합된 상품이라면 이 항목을 선택하세요."),
        ],
    },
    "s_host": {
        "tag":"🔴 예외처리 — 내재파생상품","title":"STEP 1 — 복합계약의 '주계약'이 무엇인지 확인하세요","ref":"§4.3.2 / §4.3.3","ref_key":"§4.3.2",
        "desc":"🧩 **IFRS 9 핵심 원칙 §4.3.2**: 주계약이 금융자산(채무)이면 내재파생상품을 분리하지 않고 복합계약 전체에 SPPI 테스트를 적용합니다.\n\n**핵심: 주계약이 '채권·대여금'이면 분리 금지 — 전체에 SPPI 테스트 적용**",
        "helper":{"title":"💡 실무 상품 예시","body":"**📄 금융자산 주계약 (→ 분리 금지)**\n- 전환사채 (주계약: 사채 / 내재: 전환권)\n- 신종자본증권 (주계약: 후순위채 / 내재: 이자이연)\n- COCOS (주계약: 채권 / 내재: 주식전환 트리거)\n\n**🏢 비금융자산·금융부채 주계약 (→ 분리 3요건 검토)**\n- 리스계약에 내재된 환율연계 임차료\n- 금융부채에 내재된 주가연계 이자"},
        "options":[
            ("fa_host","📄 주계약이 금융자산 (채권·대여금 등 채무상품)","복합계약의 뼈대가 이자와 원금을 받는 채권·대여금 구조라면 선택하세요. → 내재파생 분리 금지, 전체 SPPI 테스트"),
            ("other_host","🏢 주계약이 금융부채 또는 비금융자산","주계약이 금융부채·리스·상품공급계약 등이라면 선택하세요. → 분리 3요건 검토"),
        ],
    },
    "s_sep": {
        "tag":"🔴 예외처리 — 분리 3요건","title":"STEP 1-1 — 내재파생상품을 분리할 수 있는 3가지 조건을 확인하세요","ref":"§4.3.3 / §B4.3.5 / §B4.3.8","ref_key":"§4.3.3",
        "desc":"🔬 **분리 3요건:**\n① 경제적 특성·위험이 주계약과 밀접하게 관련되지 **않음**\n② 별도 계약이라면 파생상품 정의를 **충족**\n③ 복합계약 전체의 공정가치 변동을 당기손익으로 인식하지 **않음**\n\n**분리 안 하는 예외 (§B4.3.8)**: 레버리지 없는 금리캡·플로어, 인플레이션 리스료, 단위연계특성",
        "helper":None,
        "options":[
            ("sep_ok","✅ 3요건 모두 충족 — 내재파생상품을 분리할 수 있음","위 3가지 조건이 모두 해당된다면 선택하세요. → 내재파생 FVPL, 주계약 별도 처리"),
            ("sep_fail","❌ 3요건 미충족 또는 분리·측정이 실질적으로 불가능","3가지 중 하나라도 충족하지 못하거나 측정 어렵다면 선택하세요. → 복합계약 전체 FVPL"),
        ],
    },
    "s_sppi1": {
        "tag":"🔵 SPPI 테스트 ①","title":"STEP 2-① — 이자·원금이 '이상한 변수'에 연동되어 있나요?","ref":"§B4.1.7A / §B4.1.8A / §B4.1.9","ref_key":"§B4.1.7A",
        "desc":"💧 **이 단계는 '이자가 순수한 대여 비용인지, 주식·원자재 같은 딴 것과 섞인 건지' 확인하는 단계입니다.**\n\nIFRS 9는 이자가 오직 **화폐의 시간가치 + 신용위험 + 기본대여원가 + 이윤**으로만 구성되어야 한다고 봅니다.\n\n*(영향이 매우 미미한 de minimis 특성은 무시 가능 — §B4.1.18)*",
        "helper":None,
        "options":[
            ("fail_equity","❌ 주가·주가지수 연동 이자 또는 원금","이자율이나 상환금액이 코스피·S&P500 등 주가지수나 주식 가격에 따라 달라진다면 선택하세요. (예: 전환사채, ELS 연계채권)"),
            ("fail_commodity","❌ 원자재·탄소가격지수(시장추적형) 연동","금·원유·탄소가격지수 등 원자재 시장가격이 이자율 결정 기준이 된다면 선택하세요."),
            ("fail_profit","❌ 채무자 순이익·매출 비율 연동 이자","이자가 발행회사의 순이익이나 매출의 일정 %로 결정된다면 선택하세요. (예: 이익참가사채)"),
            ("fail_inverse","❌ 역변동금리 — 시장금리 오르면 이자가 오히려 내려가는 구조","시장금리 상승 시 이자가 감소하는 역방향 구조라면 선택하세요. (예: Inverse Floater)"),
            ("fail_leverage","❌ 레버리지 내재 — 독립 옵션·선도·스왑 수준의 변동성","단독 거래 시 파생상품으로 분류될 수준의 레버리지를 포함하고 있다면 선택하세요."),
            ("fail_defer","❌ 이자이연 가능 + 이연이자에 복리 미발생","발행자가 이자 지급을 미룰 수 있는데 그 밀린 이자에 복리가 붙지 않는다면 선택하세요. (예: AT1 신종자본증권 일부)"),
            ("none","✅ 위 항목 모두 해당 없음 — 순수 원리금 구조","이자가 시장금리(SOFR·EURIBOR 등)나 고정금리에만 연동되고 위 항목이 해당되지 않는다면 선택하세요. → 다음 단계 계속"),
        ],
    },
    "s_sppi2": {
        "tag":"🔵 SPPI 테스트 ②","title":"STEP 2-② — 이자율 재설정 방식이 '시간 흐름'과 맞게 설계되어 있나요?","ref":"§B4.1.9B~9D","ref_key":"§B4.1.9C",
        "desc":"⏱️ **이 단계는 '이자율 단위와 지급 주기가 같은 짝끼리 맞는지' 확인하는 단계입니다.**",
        "helper":{"title":"💡 TVM 판단 도우미 — 이자율 기간 일치 확인","body":"**핵심 질문: '이자율 재설정 주기'와 '해당 이자율의 만기 기간'이 일치합니까?**\n\n| 재설정 주기 | 사용하는 이자율 기간 | 일치 여부 |\n|---|---|---|\n| 매월 | 1개월 이자율 (1M SOFR) | ✅ 일치 |\n| 분기 | 3개월 이자율 (3M EURIBOR) | ✅ 일치 |\n| 매월 | **1년** 이자율 | ❌ 불일치 |\n| 6개월 | **5년** 이자율 (만기=5년) | ❌ 불일치 |\n\n불일치 → 벤치마크 현금흐름과 유의적 차이 여부 추가 평가 필요\n**규제이자율 예외 §B4.1.9E**: 중앙은행·금감원 규제금리가 TVM을 대략 반영한다면 허용"},
        "options":[
            ("tvm_ok","✅ 이자기산기간과 이자율 기간이 일치 (또는 영향이 미미)","'월 재설정에 1개월 금리'처럼 기간이 맞거나, 불일치 영향이 거의 없다면 선택하세요."),
            ("tvm_modified_minor","⚠️ 기간 불일치가 있으나 벤치마크와 유의적 차이 없음 확인됨","기간이 약간 맞지 않지만 벤치마크 비교 결과 유의적 차이가 없다고 판단되면 선택하세요."),
            ("tvm_fail","❌ TVM 변형이 유의적 — 벤치마크 현금흐름과 유의적 차이 확인됨","이자율 기간 불일치가 명백하고 벤치마크 현금흐름과 유의적 차이가 발생한다면 선택하세요."),
        ],
    },
    "s_sppi3": {
        "tag":"🔵 SPPI 테스트 ③","title":"STEP 2-③ — 계약에 '원리금 이외의 현금흐름'을 만드는 특수 조건이 있나요?","ref":"§B4.1.10 / §B4.1.11 / §B4.1.12","ref_key":"§B4.1.7A",
        "desc":"📋 **이 단계는 '계약서의 특수 조항이 원리금 지급 흐름을 방해하는지' 확인하는 단계입니다.**",
        "helper":None,
        "options":[
            ("clause_ok","✅ 특수 조건 없음, 또는 SPPI를 깨지 않는 조건만 존재","단순 고정·변동금리 채권이거나 중도상환 금액이 '미지급 원리금 + 합리적 보상' 수준이라면 선택하세요."),
            ("clause_exception","⚠️ 중도상환 조건이 SPPI 기준을 약간 벗어나나 §B4.1.12 예외에 해당","할인·할증 발행 채권인데 중도상환 금액이 액면+미지급이자 수준이고 초기 공정가치가 매우 작다면 선택하세요."),
            ("clause_fail","❌ 원리금 지급과 일치하지 않는 현금흐름을 만드는 계약조건 존재","주가지수 도달 시 이자율이 갑자기 뛰거나, 자산 성과에 따라 상환금이 달라지는 조건이 있다면 선택하세요."),
        ],
    },
    "s_sppi4": {
        "tag":"🔵 SPPI 테스트 ④","title":"STEP 2-④ — 이 자산이 ABS·CLO처럼 다른 자산 묶음에 연결된 구조입니까?","ref":"§B4.1.20~26","ref_key":"§B4.1.21",
        "desc":"🏗️ **이 단계는 'ABS·CLO·MBS처럼 Pool 구조 안의 한 조각인지' 확인하는 단계입니다.**\n\n**Look-through 3조건 §B4.1.21**: ①트랑슈 자체 SPPI ②기초집합 SPPI 특성 ③신용위험 노출도 ≤ 기초집합",
        "helper":{"title":"💡 트랑슈 구조 실무 상품 예시","body":"| 상품 유형 | 설명 | Look-through 필요 |\n|---|---|---|\n| ABS (자산유동화증권) | 대출채권·매출채권 유동화 | ✅ 필요 |\n| MBS (주택저당증권) | 주택담보대출 기초 발행 | ✅ 필요 |\n| CLO (대출채권담보부증권) | 기업대출 Pool 선·후순위 분리 | ✅ 필요 |\n| 일반 회사채 | 단일 발행자의 직접 채무 | ❌ 해당 없음 |"},
        "options":[
            ("tranche_no","✅ 트랑슈 구조 아님 — 단일 발행자의 일반 채무상품","일반 회사채·국채·대여금처럼 특정 발행자에게 직접 빌려주는 구조라면 선택하세요. → SPPI 충족"),
            ("tranche_pass","⚠️ 트랑슈 구조이지만 Look-through 3조건 모두 충족","ABS·CLO 등이지만 기초집합 분석 결과 3가지 조건이 모두 확인되었다면 선택하세요."),
            ("tranche_fail","❌ 트랑슈 구조이며 조건 불충족 또는 최초 인식시점에 평가 자체가 불가","ABS·CLO 등인데 기초집합이 불투명하거나 3조건 중 하나라도 충족 안 된다면 선택하세요. → FVPL"),
        ],
    },
    "s_bm": {
        "tag":"🟢 사업모형 테스트","title":"STEP 3 — 이 금융자산을 '어떤 목적으로 운용'하고 있나요?","ref":"§4.1.1⑴ / §B4.1.2 / §B4.1.2B","ref_key":"§B4.1.2B",
        "desc":"🎯 **이 단계는 '투자 목적이 이자 수취인지, 시세차익인지, 아니면 둘 다인지' 구분하는 단계입니다.**\n\n**중요**: 이 판단은 개별 상품이 아니라 **포트폴리오(집합) 수준**에서, **실제 관리 방식(사실)**에 근거해야 합니다.",
        "helper":{"title":"💡 사업모형 판단 체크리스트","body":"| 확인 항목 | AC 신호 | FVOCI 신호 | FVPL 신호 |\n|---|---|---|---|\n| 내부 성과 보고 기준 | 이자수익·ECL | 총수익(이자+매도익) | 공정가치 손익 |\n| 경영진 보상 기준 | 이자수익 달성 | 총수익률 | 공정가치 수익률 |\n| 매도 빈도·이유 | 신용위험 증가 시만 | 유동성·만기 관리 | 적극적·빈번한 매도 |\n\n**대표 포트폴리오 예시**\n- 🏦 **AC**: 은행 기업대출, 보험사 만기보유 채권\n- ⚖️ **FVOCI**: 은행 유동성 포트폴리오(LCR), ALM 채권\n- 📊 **FVPL**: IB 트레이딩 북, 헤지펀드 채권 포지션"},
        "options":[
            ("hold","🏦 계약상 현금흐름 수취 (AC 모형)","이 채권을 만기까지 보유하면서 이자와 원금을 받는 것이 핵심 목적이라면 선택하세요."),
            ("both","⚖️ 현금흐름 수취 AND 매도 둘 다 필수 (FVOCI 모형)","이자를 받으면서도 유동성 확보·만기 조절을 위해 정기적으로 매도가 반드시 필요하다면 선택하세요."),
            ("trading","📊 공정가치 실현·단기매매가 주된 목적 (FVPL 잔여범주)","매도를 통한 시세차익이 주된 수익원이거나 경영진에게 공정가치 기준으로 성과를 보고한다면 선택하세요."),
            ("ambiguous","❓ 위 중 명확하게 해당하는 것이 없음 → 추가 검토 필요","내부 보고 방식이나 운용 목적이 명확하지 않아 어느 모형인지 확신하기 어렵다면 선택하세요."),
        ],
    },
    "s_fvo": {
        "tag":"🟡 FVO 최종 확인","title":"STEP 4 — 회계 장부 불일치 해소를 위해 FVPL로 직접 지정하시겠습니까?","ref":"§4.1.5 / §B4.1.29~32","ref_key":"§4.1.5",
        "desc":"🔧 **이 단계는 '장부를 통일시키기 위한 마지막 선택지'입니다.**\n\n⚠️ **한 번 지정하면 영구적으로 취소 불가(irrevocable)합니다.**",
        "helper":None,
        "options":[
            ("fvo_yes","📌 예 — 회계불일치 해소를 위해 FVPL로 지정 (취소불가)","관련 금융부채·파생상품이 이미 FVPL로 측정되어 이 자산도 맞춰야 한다면 선택하세요."),
            ("fvo_no","✅ 아니오 — 회계불일치 없음, 원래 분류(AC/FVOCI) 확정","회계불일치가 없거나 FVO 지정이 필요하지 않다면 선택하세요."),
        ],
    },
    "s_eq_trade": {
        "tag":"🟠 지분상품","title":"STEP A-1 — 이 주식·지분을 '곧 팔 목적'으로 샀나요?","ref":"§4.1.4","ref_key":"§4.1.4",
        "desc":"⚡ **이 단계는 '단기 시세차익용인지, 장기 전략 보유인지' 구분하는 단계입니다.**",
        "helper":None,
        "options":[
            ("trade_yes","📊 예 — 단기간 내 매도를 목적으로 취득한 단기매매 주식","트레이딩 북에 편입된 주식이거나 단기 시세차익을 노리고 취득했다면 선택하세요. → FVPL 강제"),
            ("trade_no","📌 아니오 — 전략적 지분투자, 관계사 출자, 장기 보유 목적","자회사 출자금·관계사 지분·장기 전략 투자처럼 매도가 주목적이 아니라면 선택하세요. → FVOCI 선택권 검토"),
        ],
    },
    "s_eq_fvoci": {
        "tag":"🟠 지분상품","title":"STEP A-2 — OCI(기타포괄손익)에 평가 손익을 넣는 선택을 하시겠습니까?","ref":"§4.1.4 / §5.7.5~5.7.6","ref_key":"§4.1.4",
        "desc":"📌 **이 선택은 최초 인식시점에만 가능하며, 한 번 선택하면 취소할 수 없습니다(irrevocable).**\n\n⚠️ **3가지 중요한 제약:**\n1. 최초 인식시점에만 가능, 취소 불가\n2. 처분 시 OCI 손익이 P&L로 이전되지 않음(Recycling 금지)\n3. ECL 손상을 인식하지 않음",
        "helper":None,
        "options":[
            ("fvoci_yes","✅ 예 — FVOCI 지정 선택 (취소불가 / Recycling 없음 / ECL 미적용)","P&L 변동성을 낮추고 위 3가지 제약을 감수할 수 있다면 선택하세요."),
            ("fvoci_no","📊 아니오 — FVPL 유지 (기본값, 모든 변동 P&L 반영)","FVOCI 선택을 하지 않거나 요건을 재검토해야 한다면 선택하세요."),
        ],
    },
}

# 단계 레이블 한국어 매핑
_STEP_LABEL = {
    "s_asset":"STEP 0 — 자산 성격","s_host":"STEP 1 — 주계약 성격","s_sep":"STEP 1-1 — 분리 3요건",
    "s_sppi1":"STEP 2-① — 비대여변수/레버리지","s_sppi2":"STEP 2-② — TVM 변형",
    "s_sppi3":"STEP 2-③ — 계약조건","s_sppi4":"STEP 2-④ — 트랑슈",
    "s_bm":"STEP 3 — 사업모형","s_fvo":"STEP 4 — FVO 확인",
    "s_eq_trade":"STEP A-1 — 지분 단기매매","s_eq_fvoci":"STEP A-2 — FVOCI 선택권",
}
_VAL_KO = {
    "hybrid":"복합계약","equity":"지분상품","deriv":"독립 파생상품","debt":"채무상품",
    "fa_host":"금융자산 주계약","other_host":"비금융자산·금융부채 주계약",
    "sep_ok":"3요건 충족","sep_fail":"3요건 미충족",
    "fail_equity":"주가지수 연동","fail_commodity":"원자재 연동","fail_profit":"채무자 수익 연동",
    "fail_inverse":"역변동금리","fail_leverage":"레버리지 내재","fail_defer":"이자이연+복리미발생",
    "none":"비대여변수 없음","tvm_ok":"TVM 일치","tvm_modified_minor":"TVM 변형 미미",
    "tvm_fail":"TVM 변형 유의적","clause_ok":"정상 계약조건","clause_exception":"§B4.1.12 예외",
    "clause_fail":"원리금 불일치","tranche_no":"트랑슈 아님","tranche_pass":"트랑슈 3조건 충족",
    "tranche_fail":"트랑슈 조건 불충족","hold":"현금흐름 수취(AC)","both":"수취+매도(FVOCI)",
    "trading":"공정가치/단기매매(FVPL)","ambiguous":"모호함","fvo_yes":"FVO 지정",
    "fvo_no":"FVO 미지정","trade_yes":"단기매매","trade_no":"장기보유",
    "fvoci_yes":"FVOCI 지정","fvoci_no":"FVPL 유지",
}


# ══════════════════════════════════════════════════════════════════════════════
# 5. CSS 전역 주입
# ══════════════════════════════════════════════════════════════════════════════

def _inject_css():
    st.markdown(f"""
    <style>
    /* ── Pretendard 폰트 로드 ── */
    @import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css');

    /* ── 전역 폰트 적용 ── */
    html, body, [class*="css"], .stMarkdown, .stText, button, input, select, textarea {{
        font-family: 'Pretendard', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif !important;
    }}

    /* ── 메인 타이틀 ── */
    h1 {{
        color: {NAVY} !important;
        font-weight: 700 !important;
        letter-spacing: -0.03em !important;
        font-size: 1.9rem !important;
    }}

    /* ── 서브 헤더 ── */
    h2, h3 {{
        color: #1E293B !important;
        font-weight: 500 !important;
        letter-spacing: -0.02em !important;
    }}

    /* ── 버튼 기본 스타일 ── */
    .stButton > button {{
        border-radius: 8px !important;
        border: 1.5px solid #CBD5E1 !important;
        background: #F8FAFC !important;
        color: #1E293B !important;
        width: 100% !important;
        text-align: left !important;
        padding: 0.65rem 1rem !important;
        font-size: 0.9rem !important;
        line-height: 1.6 !important;
        font-weight: 400 !important;
        transition: all 0.18s ease !important;
        white-space: normal !important;
        box-shadow: 0 1px 2px rgba(0,0,0,0.04) !important;
    }}
    .stButton > button:hover {{
        border-color: {NAVY} !important;
        background: {NAVY_LIGHT} !important;
        color: {NAVY} !important;
        box-shadow: 0 2px 6px rgba(30,58,138,0.12) !important;
        transform: translateY(-1px) !important;
    }}

    /* ── Primary 버튼 ── */
    .stButton > button[kind="primary"] {{
        background: {NAVY} !important;
        color: #fff !important;
        border-color: {NAVY} !important;
        font-weight: 600 !important;
        box-shadow: 0 2px 8px rgba(30,58,138,0.25) !important;
    }}
    .stButton > button[kind="primary"]:hover {{
        background: #1E40AF !important;
        border-color: #1E40AF !important;
        box-shadow: 0 4px 12px rgba(30,58,138,0.35) !important;
    }}

    /* ── 표(Table) 스타일 ── */
    .dataframe, table {{
        background: #FFFFFF !important;
        border: 1px solid #E2E8F0 !important;
        border-radius: 8px !important;
        overflow: hidden !important;
    }}
    .dataframe th, thead th {{
        background: {NAVY_LIGHT} !important;
        color: {NAVY} !important;
        font-weight: 600 !important;
        border-bottom: 1px solid {NAVY_MID} !important;
        padding: 8px 12px !important;
    }}
    .dataframe td, tbody td {{
        border-bottom: 1px solid #E2E8F0 !important;
        padding: 7px 12px !important;
        font-size: 0.88rem !important;
        color: #334155 !important;
    }}
    .dataframe tr:last-child td {{ border-bottom: none !important; }}
    .dataframe tr:hover td {{ background: #F1F5F9 !important; }}

    /* ── 사이드바 ── */
    [data-testid="stSidebar"] {{
        background: linear-gradient(180deg, {NAVY} 0%, #1E40AF 100%) !important;
    }}
    [data-testid="stSidebar"] * {{ color: #EFF6FF !important; }}
    [data-testid="stSidebar"] .stButton > button {{
        background: rgba(255,255,255,0.1) !important;
        border-color: rgba(255,255,255,0.25) !important;
        color: #EFF6FF !important;
    }}
    [data-testid="stSidebar"] .stButton > button:hover {{
        background: rgba(255,255,255,0.2) !important;
        border-color: rgba(255,255,255,0.45) !important;
    }}

    /* ── 탭 ── */
    .stTabs [data-baseweb="tab-list"] {{
        background: rgba(255,255,255,0.08) !important;
        border-radius: 8px !important;
        padding: 3px !important;
    }}
    .stTabs [data-baseweb="tab"] {{
        color: #93C5FD !important;
        font-size: 0.8rem !important;
        font-weight: 500 !important;
        border-radius: 6px !important;
    }}
    .stTabs [aria-selected="true"] {{
        background: rgba(255,255,255,0.2) !important;
        color: #FFFFFF !important;
    }}

    /* ── 진행 바 ── */
    .stProgress > div > div > div > div {{
        background: {NAVY} !important;
        border-radius: 4px !important;
    }}

    /* ── 메트릭 카드 ── */
    [data-testid="metric-container"] {{
        background: #FFFFFF !important;
        border: 1px solid #E2E8F0 !important;
        border-radius: 12px !important;
        padding: 1rem !important;
        box-shadow: 0 1px 4px rgba(0,0,0,0.06) !important;
    }}

    /* ── expander ── */
    .streamlit-expanderHeader {{
        background: {NAVY_LIGHT} !important;
        border-radius: 6px !important;
        font-weight: 500 !important;
        color: {NAVY} !important;
    }}

    /* ── 코드 블록 ── */
    code {{
        background: {NAVY_LIGHT} !important;
        color: {NAVY} !important;
        border-radius: 4px !important;
        padding: 1px 5px !important;
        font-size: 0.85em !important;
    }}

    /* ── 분류 태그 뱃지 ── */
    .tag-badge {{
        display: inline-block;
        padding: 3px 12px;
        border-radius: 20px;
        font-size: 0.75rem;
        font-weight: 600;
        letter-spacing: 0.01em;
        margin-bottom: 0.5rem;
    }}

    /* ════════════════════════════════════════════════════════
       파일 업로더 — 사이드바 와일드카드(*) 덮어쓰기 수정
       [data-testid="stSidebar"] * {{ color: #EFF6FF }} 규칙보다
       더 구체적인 선택자로 우선순위를 강제 적용합니다.
       ════════════════════════════════════════════════════════ */

    /* 전체 카드 컨테이너: 흰색 배경 + 부드러운 테두리 */
    [data-testid="stSidebar"] [data-testid="stFileUploader"] {{
        background: #FFFFFF !important;
        border-radius: 12px !important;
        padding: 8px !important;
        box-shadow: 0 1px 3px rgba(0, 0, 0, 0.15) !important;
    }}

    /* 드롭존: 흰색 바탕 + 연한 회색 점선 테두리 */
    [data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"] {{
        background-color: #FFFFFF !important;
        border: 2px dashed #CBD5E1 !important;
        border-radius: 10px !important;
        padding: 0.75rem !important;
        transition: border-color 0.2s ease, background-color 0.2s ease !important;
    }}
    [data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"]:hover {{
        border-color: #1E3A8A !important;
        background-color: #EFF6FF !important;
    }}

    /* 드롭존 안내 문구: 다크 그레이 */
    [data-testid="stSidebar"] [data-testid="stFileUploaderDropzoneInstructions"] span,
    [data-testid="stSidebar"] [data-testid="stFileUploaderDropzoneInstructions"] small,
    [data-testid="stSidebar"] [data-testid="stFileUploaderDropzoneInstructions"] p,
    [data-testid="stSidebar"] [data-testid="stFileUploaderDropzoneInstructions"] div {{
        color: #475569 !important;
        font-size: 0.8rem !important;
    }}

    /* ── Browse files 버튼: 네이비 블루 배경 + 흰색 글씨 ──
       와일드카드 규칙을 이기려면 선택자 구체성을 최대한 높입니다. */
    [data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"] button {{
        background-color: #1E3A8A !important;
        color: #FFFFFF !important;
        border: none !important;
        border-radius: 8px !important;
        padding: 0.38rem 1rem !important;
        font-size: 0.82rem !important;
        font-weight: 600 !important;
        letter-spacing: 0.01em !important;
        cursor: pointer !important;
        transition: background-color 0.2s ease, box-shadow 0.2s ease !important;
        box-shadow: 0 1px 3px rgba(30, 58, 138, 0.3) !important;
        width: auto !important;
        text-align: center !important;
    }}
    [data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"] button:hover {{
        background-color: #3B82F6 !important;
        color: #FFFFFF !important;
        box-shadow: 0 3px 8px rgba(59, 130, 246, 0.4) !important;
    }}
    /* 버튼 내부 텍스트 span도 별도로 덮어씌움 */
    [data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"] button span,
    [data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"] button p {{
        color: #FFFFFF !important;
    }}

    /* 업로드된 파일명: 다크 그레이 + 연한 배경 카드 */
    [data-testid="stSidebar"] [data-testid="stFileUploaderFileName"],
    [data-testid="stSidebar"] [data-testid="stFileUploaderFile"] span,
    [data-testid="stSidebar"] [data-testid="stFileUploaderFile"] div {{
        color: #1E293B !important;
        font-size: 0.82rem !important;
        font-weight: 500 !important;
    }}
    [data-testid="stSidebar"] [data-testid="stFileUploaderFile"] {{
        background: #F8FAFC !important;
        border: 1px solid #E2E8F0 !important;
        border-radius: 8px !important;
        padding: 0.35rem 0.7rem !important;
        margin-top: 0.4rem !important;
    }}

    /* 파일 삭제(×) 버튼 */
    [data-testid="stSidebar"] [data-testid="stFileUploaderFile"] button {{
        background: transparent !important;
        border: none !important;
        color: #94A3B8 !important;
        width: auto !important;
        padding: 0.15rem 0.3rem !important;
        box-shadow: none !important;
    }}
    [data-testid="stSidebar"] [data-testid="stFileUploaderFile"] button:hover {{
        background: #FEF2F2 !important;
        color: #DC2626 !important;
    }}
    [data-testid="stSidebar"] [data-testid="stFileUploaderFile"] button span {{
        color: inherit !important;
    }}
    </style>
    """, unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# 6. 세션 상태 초기화 (전수 점검)
# ══════════════════════════════════════════════════════════════════════════════

def _init_session():
    """누락 없이 모든 session_state 키를 초기화"""
    defaults = {
        "answers": {},
        "history": [],
        "show_result": False,
        "ai_mode": False,
        "ai_result": None,
        "ai_overrides": {},
        "ai_confirmed": False,
        "contract_text_preview": "",
        "wizard_started": False,
        # v5 신규 키
        "uploaded_file_infos": [],      # extract_text_from_files() 반환값
        "bm_inference": None,           # infer_bm() 반환값
        "bm_override": None,            # 사용자가 BM 패널에서 수정한 값
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


def _full_reset():
    """전체 상태 초기화"""
    reset_map = {
        "answers": {}, "ai_overrides": {},
        "history": [], "uploaded_file_infos": [],
        "ai_result": None, "bm_inference": None, "bm_override": None,
        "show_result": False, "ai_mode": False, "ai_confirmed": False, "wizard_started": False,
        "contract_text_preview": "",
    }
    for k, v in reset_map.items():
        st.session_state[k] = v
    st.rerun()


# ══════════════════════════════════════════════════════════════════════════════
# 7. 사이드바 (탭 통합)
# ══════════════════════════════════════════════════════════════════════════════

def _render_sidebar():
    with st.sidebar:
        # 로고 / 브랜딩
        st.markdown("""
        <div style="text-align:center;padding:1rem 0 0.5rem">
            <div style="font-size:2rem">📊</div>
            <div style="font-size:1rem;font-weight:700;color:#FFFFFF;letter-spacing:-0.02em">IFRS 9 분류 마법사</div>
            <div style="font-size:0.7rem;color:#93C5FD;margin-top:2px">K-IFRS 1109호 기반 · v4</div>
        </div>
        """, unsafe_allow_html=True)
        st.divider()

        tab_doc, tab_dict = st.tabs(["📄 문서 분석", "📖 용어 사전"])

        # ── 탭 1: 문서 분석 ──────────────────────────────────────────────────
        with tab_doc:
            st.markdown("**🤖 AI 계약서 자동 분석**")
            st.caption("계약서·기안문 등 여러 파일을 한 번에 업로드하면\nAI가 STEP 0~2와 사업모형(BM)을 자동 제안합니다.")

            libs_msg = []
            if not _PDF_OK: libs_msg.append("pdfplumber")
            if not _DOCX_OK: libs_msg.append("python-docx")
            if libs_msg:
                st.warning(f"미설치: `pip install {' '.join(libs_msg)}`")

            uploaded_files = st.file_uploader(
                "계약서 / 기안문 업로드 (PDF · DOCX, 복수 선택 가능)",
                type=["pdf", "docx"],
                accept_multiple_files=True,          # ← v5: 다중 업로드
                key="file_uploader",
                help="계약서 + 내부 기안문을 함께 올리면 사업모형을 더 정확하게 추론합니다.",
                label_visibility="collapsed",
            )

            # 업로드된 파일 목록 미리보기
            if uploaded_files:
                st.markdown(
                    f'<div style="font-size:0.75rem;color:#93C5FD;margin:4px 0 6px">'
                    f'📎 {len(uploaded_files)}개 파일 선택됨</div>',
                    unsafe_allow_html=True,
                )
                for uf in uploaded_files:
                    ext_icon = "📄" if uf.name.lower().endswith(".pdf") else "📝"
                    size_kb = round(uf.size / 1024, 1) if hasattr(uf, "size") else "?"
                    st.markdown(
                        f'<div style="background:rgba(255,255,255,0.12);border-radius:6px;'
                        f'padding:4px 8px;margin-bottom:3px;font-size:0.75rem;color:#DBEAFE">'
                        f'{ext_icon} {uf.name} <span style="opacity:0.65">({size_kb} KB)</span>'
                        f'</div>',
                        unsafe_allow_html=True,
                    )

            if uploaded_files:
                if st.button("🪄 AI 분석 시작", type="primary", use_container_width=True):
                    errors, all_texts, files_info = [], [], []

                    with st.spinner(f"📄 텍스트 추출 중 ({len(uploaded_files)}개 파일)..."):
                        files_info = extract_text_from_files(uploaded_files)

                    # 오류 파일 분리
                    for fi in files_info:
                        if fi["error"]:
                            errors.append(f"⚠️ {fi['filename']}: {fi['text']}")
                        else:
                            all_texts.append(fi["text"])

                    if errors:
                        for msg in errors:
                            st.error(msg)
                    if not all_texts:
                        st.error("텍스트를 추출할 수 있는 파일이 없습니다.")
                    else:
                        # 통합 텍스트 (모든 파일 결합)
                        combined_text = "\n\n".join(all_texts)

                        with st.spinner("🔍 SPPI·자산 성격 분석 중..."):
                            result = ai_analyze(combined_text, files_info)

                        # 세션 상태 저장
                        st.session_state.ai_mode = True
                        st.session_state.ai_result = result
                        st.session_state.ai_overrides = {}
                        st.session_state.ai_confirmed = False
                        st.session_state.answers = {}
                        st.session_state.history = []
                        st.session_state.show_result = False
                        st.session_state.wizard_started = True
                        st.session_state.uploaded_file_infos = files_info
                        st.session_state.bm_inference = result.get("bm_inference")
                        st.session_state.bm_override = None
                        # 미리보기: 각 파일 앞 500자 통합
                        preview_parts = []
                        for fi in files_info:
                            if not fi["error"]:
                                preview_parts.append(
                                    f"=== [{fi['role']}] {fi['filename']} ===\n{fi['text'][:500]}"
                                )
                        st.session_state.contract_text_preview = "\n\n".join(preview_parts)[:3000]
                        st.rerun()

            st.divider()
            st.markdown("**✏️ 수동 분류**")
            if st.button("단계별 직접 분류 시작", use_container_width=True):
                st.session_state.ai_mode = False
                st.session_state.ai_result = None
                st.session_state.ai_confirmed = False
                st.session_state.answers = {}
                st.session_state.history = []
                st.session_state.show_result = False
                st.session_state.wizard_started = True
                st.session_state.uploaded_file_infos = []
                st.session_state.bm_inference = None
                st.session_state.bm_override = None
                st.rerun()

            if st.session_state.wizard_started:
                st.divider()
                if st.button("🔄 처음부터 다시", use_container_width=True):
                    _full_reset()

            # AI 완료 상태 표시
            if st.session_state.ai_mode and st.session_state.ai_result:
                file_cnt = len(st.session_state.get("uploaded_file_infos", []))
                st.success(f"✅ AI 분석 완료 ({file_cnt}개 파일)\n메인 화면에서 확인하세요")

        # ── 탭 2: 용어 사전 ──────────────────────────────────────────────────
        with tab_dict:
            st.markdown("**📖 실무 용어 사전**")
            glossary = [
                ("SPPI", "Solely Payments of Principal and Interest.\n이자가 오직 대여에 대한 대가(시간가치+신용위험+기본원가+이윤)로만 구성되어야 AC·FVOCI가 가능합니다."),
                ("사업모형 (BM)", "금융자산을 어떻게 관리하는지에 대한 경영진의 의사결정. 포트폴리오 수준에서 내부 증거(보고체계·보상·매도이력)로 판단합니다."),
                ("AC (상각후원가)", "Amortised Cost. 원금+이자를 만기까지 받는 채무상품에 적용. 유효이자율법으로 측정. ECL 손상 인식."),
                ("FVOCI", "공정가치로 측정하되, 평가손익을 OCI에 반영. 채무상품은 처분 시 Recycling, 지분상품은 Recycling 금지."),
                ("FVPL", "공정가치로 측정하고 모든 변동을 당기손익(P&L)에 반영. 파생상품·단기매매의 기본 방법."),
                ("ECL (기대신용손실)", "Expected Credit Loss. AC·채무 FVOCI 자산에만 적용. 신용위험 유의적 증가 시 전체기간 ECL."),
                ("Recycling", "OCI 누적 손익을 처분 시 P&L로 옮기는 것. 채무상품 FVOCI는 발생, 지분상품 FVOCI는 영구 금지."),
                ("TVM (화폐의 시간가치)", "Time Value of Money. 이자가 TVM을 적절히 반영해야 SPPI를 충족합니다."),
                ("내재파생상품", "주계약 안에 숨어있는 파생상품. 금융자산 주계약이면 분리 금지, 전체를 하나로 SPPI 테스트(§4.3.2)."),
                ("FVO (공정가치 지정선택권)", "회계불일치 해소를 위해 AC/FVOCI 대신 FVPL로 지정하는 취소불가 선택권(§4.1.5)."),
                ("트랑슈 / ABS·CLO", "여러 자산을 묶어 선순위·후순위로 나눈 구조화 금융상품. SPPI 판단 시 Look-through 분석 필요."),
                ("복합계약", "주계약(채권 등)과 내재파생상품이 결합된 금융상품. 예: 전환사채, 신종자본증권."),
            ]
            for term, explanation in glossary:
                with st.expander(f"**{term}**"):
                    st.markdown(explanation)


# ══════════════════════════════════════════════════════════════════════════════
# 8. 상단 진행 바 (단계 번호 + 퍼센트)
# ══════════════════════════════════════════════════════════════════════════════

def _render_progress(ans: dict):
    seq = get_step_sequence(ans)
    step_num = len(seq)
    at = ans.get("s_asset")
    if at == "deriv":    total = 1
    elif at == "equity": total = 3
    elif at == "hybrid":
        host = ans.get("s_host")
        total = 4 if host == "other_host" else 8
    else:                total = 8

    pct = min(step_num / max(total, 1), 0.97)
    pct_int = int(pct * 100)

    # 단계 레이블 표시
    step_labels_all = [
        ("S0","자산 성격"), ("S1","복합계약"), ("S2","SPPI"), ("S3","사업모형"), ("S4","확인/결과")
    ]
    current_group = (
        0 if step_num <= 1 else
        1 if ans.get("s_asset") == "hybrid" and step_num == 2 else
        2 if step_num <= 5 else
        3 if step_num == 6 else 4
    )

    cols = st.columns([3, 1])
    with cols[0]:
        st.progress(pct, text=f"**단계 {step_num}** / 최대 {total}단계 완료")
    with cols[1]:
        st.markdown(
            f'<div style="text-align:right;padding-top:0.3rem;font-size:0.85rem;color:#64748B">'
            f'<b style="color:{NAVY}">{pct_int}%</b> 완료</div>',
            unsafe_allow_html=True,
        )

    # 단계 인디케이터
    ind_cols = st.columns(5)
    for i, (code, name) in enumerate(step_labels_all):
        with ind_cols[i]:
            if i < current_group:
                bg, fg = GREEN_LIGHT, GREEN
                icon = "✓"
            elif i == current_group:
                bg, fg = NAVY_LIGHT, NAVY
                icon = "●"
            else:
                bg, fg = "#F1F5F9", "#94A3B8"
                icon = "○"
            st.markdown(
                f'<div style="text-align:center;background:{bg};border-radius:8px;padding:5px 2px;'
                f'font-size:0.72rem;color:{fg};font-weight:{"600" if i==current_group else "400"}">'
                f'{icon} {code}<br><span style="font-size:0.65rem;opacity:0.8">{name}</span></div>',
                unsafe_allow_html=True,
            )


# ══════════════════════════════════════════════════════════════════════════════
# 9. 환영 화면 (파일 업로드 전)
# ══════════════════════════════════════════════════════════════════════════════

def _render_welcome():
    st.markdown(f"""
    <div style="background:linear-gradient(135deg,{NAVY_LIGHT} 0%,{NAVY_MID} 100%);
    border-radius:16px;padding:2.5rem 2rem;margin-bottom:1.5rem;border:1px solid {NAVY_MID}">
        <h2 style="color:{NAVY};margin:0 0 0.5rem;font-size:1.5rem;font-weight:700">
            IFRS 9 금융자산 분류 마법사에 오신 것을 환영합니다 🎯
        </h2>
        <p style="color:#475569;margin:0;line-height:1.7;font-size:0.95rem">
            K-IFRS 1109호 4장 · 부록B · PwC 실무 가이드라인을 기반으로<br>
            금융자산을 <b>AC · FVOCI · FVPL</b> 로 올바르게 분류할 수 있도록 안내합니다.
        </p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2, col3 = st.columns(3)
    cards = [
        ("📄", "계약서 자동 분석", "왼쪽 사이드바에서 PDF/DOCX를 업로드하면 AI가 핵심 조항을 스캔하여 STEP 0~2를 자동 제안합니다.", NAVY_LIGHT, NAVY),
        ("🧭", "단계별 질문 안내", "각 단계마다 쉬운 비유와 가이드형 선택지로 복잡한 기준서 로직을 체계적으로 안내합니다.", "#F0FDF4", GREEN),
        ("📊", "대시보드 결과", "최종 분류 결과를 AC/FVOCI/FVPL 대시보드로 표시하고 회계처리 요약표를 제공합니다.", "#FFFBEB", "#92400E"),
    ]
    for col, (icon, title, desc, bg, color) in zip([col1, col2, col3], cards):
        with col:
            st.markdown(
                f'<div style="background:{bg};border-radius:12px;padding:1.2rem;height:160px;'
                f'border:1px solid #E2E8F0">'
                f'<div style="font-size:1.5rem;margin-bottom:0.4rem">{icon}</div>'
                f'<div style="font-weight:600;color:{color};margin-bottom:0.35rem;font-size:0.92rem">{title}</div>'
                f'<div style="font-size:0.8rem;color:#64748B;line-height:1.55">{desc}</div>'
                f'</div>',
                unsafe_allow_html=True,
            )

    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown(
        f'<div style="background:#FFF7ED;border:1px solid #FED7AA;border-radius:10px;'
        f'padding:0.8rem 1.1rem;font-size:0.85rem;color:#92400E">'
        f'💡 <b>시작 방법:</b> 왼쪽 사이드바의 <b>[문서 분석]</b> 탭에서 계약서를 업로드하거나, '
        f'<b>"단계별 직접 분류 시작"</b> 버튼을 클릭하세요.</div>',
        unsafe_allow_html=True,
    )


# ══════════════════════════════════════════════════════════════════════════════
# 10. AI 분석 결과 확인 화면
# ══════════════════════════════════════════════════════════════════════════════

def _render_ai_confirm():
    result: dict = st.session_state.ai_result
    overrides: dict = st.session_state.ai_overrides
    proposed: dict = result["proposed_answers"]
    evidence: list = result["evidence_items"]
    conflicts: list = result["conflict_flags"]
    bm_inf: dict | None = result.get("bm_inference") or st.session_state.get("bm_inference")
    files_info: list = st.session_state.get("uploaded_file_infos", [])
    file_cnt = len(files_info)

    # ── 헤더 배너 ─────────────────────────────────────────────────────────────
    st.markdown(
        f'<div style="background:{NAVY_LIGHT};border:1.5px solid {NAVY};border-radius:12px;'
        f'padding:1rem 1.4rem;margin-bottom:1.2rem">'
        f'<span style="font-size:1.1rem;font-weight:700;color:{NAVY}">🤖 AI 계약서 분석 결과'
        f'{"  ·  " + str(file_cnt) + "개 파일" if file_cnt > 1 else ""}'
        f'</span><br>'
        f'<span style="font-size:0.83rem;color:#475569">'
        f'STEP 0~2는 AI 제안을 검토·수정하세요. 사업모형(STEP 3)은 AI 추론 결과를 참고해 직접 확정하세요.'
        f'</span>'
        f'</div>',
        unsafe_allow_html=True,
    )

    # ── 분석된 파일 목록 ──────────────────────────────────────────────────────
    if file_cnt > 0:
        with st.expander(f"📎 분석된 파일 목록 ({file_cnt}개)", expanded=False):
            role_colors = {"계약서": ("#DBEAFE","#1E40AF"), "기안문": ("#D1FAE5","#065F46"), "기타": ("#F1F5F9","#475569")}
            for fi in files_info:
                bg_r, fg_r = role_colors.get(fi["role"], ("#F1F5F9","#475569"))
                status = "⚠️ 추출 오류" if fi["error"] else "✅ 추출 완료"
                st.markdown(
                    f'<div style="display:flex;align-items:center;gap:8px;padding:5px 0;border-bottom:0.5px solid #E2E8F0">'
                    f'<span style="background:{bg_r};color:{fg_r};padding:2px 8px;border-radius:4px;font-size:0.72rem;font-weight:600">{fi["role"]}</span>'
                    f'<span style="font-size:0.83rem;color:#334155;flex:1">{fi["filename"]}</span>'
                    f'<span style="font-size:0.72rem;color:#64748B">{status}</span>'
                    f'</div>',
                    unsafe_allow_html=True,
                )

    # ── 충돌 경고 ─────────────────────────────────────────────────────────────
    for cf in conflicts:
        st.warning(f"⚠️ **{_STEP_LABEL.get(cf['step_id'], cf['step_id'])}**: {cf['message']}")

    # ── STEP 0~2 분석 결과 ────────────────────────────────────────────────────
    show_steps = [e for e in evidence if e["step_id"] != "s_bm"]
    if not show_steps:
        st.info("SPPI·자산 성격 관련 키워드가 감지되지 않았습니다. 수동 분류를 진행해주세요.")
        if st.button("✏️ 수동 분류로 전환", use_container_width=True):
            st.session_state.ai_mode = False
            st.rerun()
        return

    st.markdown("#### 📋 STEP 0~2 — 자산 성격 · SPPI 분석 결과")
    for ev in show_steps:
        step_id = ev["step_id"]
        proposed_val = proposed.get(step_id, "")
        current_val = overrides.get(step_id, proposed_val)
        conf = ev["confidence"]
        conf_label, conf_bg, conf_fg = _CONF_CFG.get(conf, ("낮음","#FEE2E2","#991B1B"))

        with st.container():
            st.markdown(
                f'<div style="background:{conf_bg};border-radius:8px;padding:0.7rem 1rem;'
                f'margin-bottom:0.5rem;border:0.5px solid #DDD">'
                f'<b style="color:{conf_fg}">{_STEP_LABEL.get(step_id, step_id)}</b> '
                f'<span style="font-size:0.72rem;background:#FFF;padding:2px 7px;border-radius:4px;'
                f'margin-left:6px;color:{conf_fg}">신뢰도: {conf_label}</span>'
                f'</div>',
                unsafe_allow_html=True,
            )
            c_info, c_edit = st.columns([2, 1])
            with c_info:
                st.markdown(f"**AI 제안:** `{_VAL_KO.get(proposed_val, proposed_val)}`")
                st.markdown(f"**근거:** {ev['basis']}")
                st.caption(f"기준서: `{ev['std_ref']}` | 키워드: {', '.join(f'`{k}`' for k in ev['keywords_found'][:5])}")
            with c_edit:
                step_def = STEP_DEFS.get(step_id)
                if step_def:
                    opts = [(v, lbl) for v, lbl, _ in step_def["options"]]
                    opt_labels = [lbl for v, lbl in opts]
                    opt_vals = [v for v, lbl in opts]
                    try:
                        idx = opt_vals.index(current_val)
                    except ValueError:
                        idx = 0
                    selected = st.selectbox("✏️ 수정", options=opt_labels, index=idx,
                                            key=f"ai_ovr_{step_id}", label_visibility="collapsed")
                    new_val = opt_vals[opt_labels.index(selected)]
                    if new_val != overrides.get(step_id, proposed_val):
                        st.session_state.ai_overrides[step_id] = new_val
        st.divider()

    # ══════════════════════════════════════════════════════════════════════════
    # BM 추론 패널 (v5 신규)
    # ══════════════════════════════════════════════════════════════════════════
    st.markdown("#### 🎯 STEP 3 — 사업모형(BM) AI 추론 결과")

    _BM_KO = {
        "hold":      ("AC — 상각후원가 모형", "#D1FAE5", "#065F46"),
        "both":      ("FVOCI — 수취·매도 병행 모형", "#DBEAFE", "#1E40AF"),
        "trading":   ("FVPL — 공정가치 실현 모형", "#FEE2E2", "#991B1B"),
        "ambiguous": ("추가 검토 필요", "#FEF3C7", "#92400E"),
    }

    if bm_inf:
        bm_val = bm_inf.get("proposed_bm", "ambiguous")
        bm_conf = bm_inf.get("confidence", "low")
        bm_label, bm_bg, bm_fg = _BM_KO.get(bm_val, _BM_KO["ambiguous"])
        bm_conf_label, _, _ = _CONF_CFG.get(bm_conf, ("낮음","",""))
        ac_s  = bm_inf.get("ac_score", 0)
        fo_s  = bm_inf.get("fvoci_score", 0)
        fp_s  = bm_inf.get("fvpl_score", 0)

        # 추론 결과 배너
        st.markdown(
            f'<div style="background:{bm_bg};border:1.5px solid {bm_fg};border-radius:10px;'
            f'padding:0.9rem 1.2rem;margin-bottom:0.8rem">'
            f'<div style="font-size:1rem;font-weight:700;color:{bm_fg};margin-bottom:0.3rem">'
            f'추정 사업모형: {bm_label}'
            f'<span style="font-size:0.72rem;background:rgba(0,0,0,0.08);padding:2px 8px;'
            f'border-radius:4px;margin-left:8px">신뢰도: {bm_conf_label}</span></div>'
            f'<div style="font-size:0.83rem;color:{bm_fg};opacity:0.9">{bm_inf.get("summary","")}</div>'
            f'</div>',
            unsafe_allow_html=True,
        )

        # 점수 막대
        sc_cols = st.columns(3)
        total_s = max(ac_s + fo_s + fp_s, 1)
        for col, (label, score, color) in zip(sc_cols, [
            ("AC 신호", ac_s, "#059669"), ("FVOCI 신호", fo_s, "#3B82F6"), ("FVPL 신호", fp_s, "#EF4444")
        ]):
            pct_s = int(score / total_s * 100)
            with col:
                st.markdown(
                    f'<div style="text-align:center;font-size:0.75rem;color:#64748B;margin-bottom:3px">'
                    f'{label}</div>'
                    f'<div style="background:#E2E8F0;border-radius:4px;height:8px">'
                    f'<div style="background:{color};width:{pct_s}%;height:8px;border-radius:4px"></div></div>'
                    f'<div style="text-align:center;font-size:0.72rem;color:{color};font-weight:600;margin-top:2px">'
                    f'{score:.1f}점 ({pct_s}%)</div>',
                    unsafe_allow_html=True,
                )

        # 근거 문구 (감지된 문장)
        ev_lines = bm_inf.get("evidence_lines", [])
        if ev_lines:
            with st.expander("📄 판단 근거 — 감지된 문구 상세", expanded=True):
                sig_colors = {"AC": ("#D1FAE5","#065F46"), "FVOCI": ("#DBEAFE","#1E40AF"), "FVPL": ("#FEE2E2","#991B1B")}
                for e in ev_lines:
                    sbg, sfg = sig_colors.get(e["signal"], ("#F1F5F9","#334155"))
                    st.markdown(
                        f'<div style="border-left:3px solid {sfg};padding:5px 10px;margin-bottom:5px;'
                        f'background:{sbg};border-radius:0 6px 6px 0">'
                        f'<span style="font-size:0.7rem;font-weight:600;color:{sfg}">'
                        f'{e["signal"]} 신호 · [{e["role"]}] {e["source"]}</span><br>'
                        f'<span style="font-size:0.8rem;color:#334155">키워드: <b>{e["keyword"]}</b></span><br>'
                        f'<span style="font-size:0.78rem;color:#64748B;font-style:italic">"{e["text"]}"</span>'
                        f'</div>',
                        unsafe_allow_html=True,
                    )

        # 사용자 수정 selectbox
        bm_opts = [("hold","🏦 AC — 계약상 현금흐름 수취"),
                   ("both","⚖️ FVOCI — 수취+매도 병행"),
                   ("trading","📊 FVPL — 공정가치 실현·단기매매"),
                   ("ambiguous","❓ 추가 검토 필요")]
        bm_opt_labels = [lbl for v, lbl in bm_opts]
        bm_opt_vals   = [v for v, lbl in bm_opts]
        cur_bm = st.session_state.get("bm_override") or bm_val
        try:
            bm_idx = bm_opt_vals.index(cur_bm)
        except ValueError:
            bm_idx = 0
        st.caption("사업모형을 수정하려면 아래 선택지를 변경하세요.")
        bm_selected = st.selectbox(
            "사업모형 최종 선택",
            options=bm_opt_labels,
            index=bm_idx,
            key="bm_select_override",
        )
        bm_new_val = bm_opt_vals[bm_opt_labels.index(bm_selected)]
        if bm_new_val != st.session_state.get("bm_override"):
            st.session_state.bm_override = bm_new_val

    else:
        st.info("업로드된 문서에서 사업모형 관련 키워드가 감지되지 않았습니다. 아래에서 직접 선택해주세요.")
        bm_opts_manual = [("hold","🏦 AC — 계약상 현금흐름 수취"),
                          ("both","⚖️ FVOCI — 수취+매도 병행"),
                          ("trading","📊 FVPL — 공정가치 실현·단기매매"),
                          ("ambiguous","❓ 추가 검토 필요")]
        bm_opt_labels_m = [lbl for v, lbl in bm_opts_manual]
        bm_opt_vals_m   = [v for v, lbl in bm_opts_manual]
        bm_sel_m = st.selectbox("사업모형 선택", options=bm_opt_labels_m, key="bm_select_manual")
        st.session_state.bm_override = bm_opt_vals_m[bm_opt_labels_m.index(bm_sel_m)]

    # ── 텍스트 미리보기 ───────────────────────────────────────────────────────
    if st.session_state.contract_text_preview:
        with st.expander("📄 추출된 텍스트 미리보기", expanded=False):
            st.text_area("", value=st.session_state.contract_text_preview,
                         height=180, disabled=True, label_visibility="collapsed")

    st.markdown("---")

    # 신뢰도에 따른 안내 메시지
    if bm_inf:
        conf_now = bm_inf.get("confidence", "low")
        sppi_note = bm_inf.get("sppi_attention_note")
        if conf_now == "high":
            st.success("✅ AI 분석 신뢰도 **높음** — 사업모형 선택 단계를 건너뛰고 바로 결과로 이동합니다.", icon="✅")
        elif conf_now == "medium":
            st.info("🟡 AI 분석 신뢰도 **보통** — 사업모형 AI 추론값을 자동 반영하여 결과로 이동합니다.", icon="ℹ️")
        else:
            st.warning("🔴 AI 분석 신뢰도 **낮음** — 추론된 최유력 사업모형을 추천값으로 세팅합니다. 결과 화면에서 수정 가능합니다.", icon="⚠️")
        if sppi_note:
            st.warning(sppi_note, icon="🔍")
    else:
        st.info("**확인 버튼을 누르면 AI 분석 결과(STEP 0~2)와 사업모형이 자동 저장되어 결과로 이동합니다.**", icon="ℹ️")

    c1, c2 = st.columns([2, 1])
    with c1:
        if st.button("🚀 확인 및 분류 결과 보기", type="primary", use_container_width=True):
            final = {}
            for ev in show_steps:
                sid = ev["step_id"]
                final[sid] = st.session_state.ai_overrides.get(sid, proposed.get(sid, ""))

            # BM 자동 세팅 — 수동 오버라이드 > proposed_bm > suggested_bm > ambiguous
            final_bm = (
                st.session_state.get("bm_override")
                or (bm_inf.get("proposed_bm") if bm_inf else None)
                or (bm_inf.get("suggested_bm") if bm_inf else None)
                or "ambiguous"
            )
            final["s_bm"] = final_bm

            # s_bm 스킵: history에 추가하여 단계 자동 통과
            history_keys = list(final.keys())
            if "s_bm" not in history_keys:
                history_keys.append("s_bm")

            # hold·both → FVO 기본값 자동 세팅
            if final_bm in ("hold", "both"):
                final["s_fvo"] = "fvo_no"
                if "s_fvo" not in history_keys:
                    history_keys.append("s_fvo")

            # RCPS·지분 감지 시 hybrid 교정
            is_eq = bm_inf.get("is_equity_signal", False) if bm_inf else False
            eq_kws = bm_inf.get("equity_keywords", []) if bm_inf else []
            rcps_found = any(kw.lower() in (k.lower() for k in eq_kws) for kw in _RCPS_KEYWORDS)
            if is_eq and rcps_found:
                if final.get("s_asset") == "debt":
                    final["s_asset"] = "hybrid"
                st.session_state["equity_rcps_hint"] = True
                st.session_state["equity_rcps_keywords"] = eq_kws
            else:
                st.session_state["equity_rcps_hint"] = False

            # sppi_attention_note 세션 저장
            st.session_state["sppi_attention_note"] = (
                bm_inf.get("sppi_attention_note") if bm_inf else None
            )

            st.session_state.answers = final
            st.session_state.ai_confirmed = True
            st.session_state.history = history_keys
            st.session_state.show_result = True
            st.rerun()
    with c2:
        if st.button("✏️ 수동 분류로 전환", use_container_width=True):
            st.session_state.ai_mode = False
            st.session_state.answers = {}
            st.rerun()


# ══════════════════════════════════════════════════════════════════════════════
# 11. 단계 화면 렌더링
# ══════════════════════════════════════════════════════════════════════════════

def _render_step(current_step_id: str, ans: dict):
    step = STEP_DEFS.get(current_step_id)
    if not step:
        st.error("알 수 없는 단계입니다.")
        _full_reset()
        return

    # 태그 뱃지
    if step["tag"]:
        tag_colors = {
            "🔴":"#FEF2F2,#991B1B","🔵":"#EFF6FF,#1E40AF",
            "🟢":"#F0FDF4,#065F46","🟡":"#FFFBEB,#92400E","🟠":"#FFF7ED,#9A3412",
        }
        first_char = step["tag"][:2]
        colors = tag_colors.get(first_char, f"{NAVY_LIGHT},{NAVY}")
        bg_c, fg_c = colors.split(",")
        st.markdown(
            f'<span class="tag-badge" style="background:{bg_c};color:{fg_c}">{step["tag"]}</span>',
            unsafe_allow_html=True,
        )

    # 제목
    st.subheader(step["title"])

    # 기준서 expander
    ref_text = STD_TEXTS.get(step.get("ref_key",""))
    if ref_text:
        with st.expander(f"📘 기준서 핵심 문구 — `{step.get('ref','')}`", expanded=False):
            st.info(ref_text)
    else:
        st.caption(f"📌 기준서: `{step.get('ref','')}`")

    st.markdown(step["desc"])

    # Helper
    if step["helper"]:
        with st.expander(step["helper"]["title"], expanded=True):
            st.markdown(step["helper"]["body"])

    st.write("")

    # 선택지 버튼
    chosen = ans.get(current_step_id)
    for val, label, sub in step["options"]:
        is_sel = chosen == val
        full_label = f"{'✔ ' if is_sel else ''}{label}\n\n*{sub}*"
        if st.button(full_label, key=f"btn_{current_step_id}_{val}",
                     type="primary" if is_sel else "secondary"):
            _pick(current_step_id, val)

    st.divider()

    # 네비게이션
    c_back, c_next, c_reset = st.columns([1, 1, 1])
    with c_back:
        if st.session_state.history:
            if st.button("← 이전 단계", use_container_width=True):
                _go_back()
    with c_next:
        if chosen:
            if st.button("다음 →", type="primary", use_container_width=True):
                _go_next()
        else:
            st.button("다음 →", disabled=True, use_container_width=True,
                      help="먼저 위의 선택지 중 하나를 클릭하세요")
    with c_reset:
        if st.button("🔄 처음부터", use_container_width=True):
            _full_reset()

    # 입력 경로 미리보기
    if ans:
        with st.expander("🗺️ 지금까지 입력한 경로 확인", expanded=False):
            for k, v in ans.items():
                st.code(f"{_STEP_LABEL.get(k, k)}: {_VAL_KO.get(v, v)}", language=None)


# ══════════════════════════════════════════════════════════════════════════════
# 12. 사업모형 단계 (AI 모드 후)
# ══════════════════════════════════════════════════════════════════════════════

def _render_bm_after_ai():
    st.success("✅ **AI 분석 완료 — STEP 0~2가 자동으로 설정되었습니다.** 아래에서 사업모형(STEP 3)만 직접 선택하면 최종 분류 결과가 나옵니다.")
    with st.expander("🤖 AI가 설정한 이전 단계 요약", expanded=True):
        for k, v in st.session_state.answers.items():
            st.markdown(f"- **{_STEP_LABEL.get(k, k)}**: `{_VAL_KO.get(v, v)}`")
    st.divider()
    _render_step("s_bm", st.session_state.answers)


# ══════════════════════════════════════════════════════════════════════════════
# 13. 결과 대시보드
# ══════════════════════════════════════════════════════════════════════════════

def _render_result(ans: dict):
    r = compute_result(ans)

    color_cfg = {
        "green":  ("#D1FAE5","#065F46","#059669","✅"),
        "blue":   ("#DBEAFE","#1E40AF","#3B82F6","🔵"),
        "orange": ("#FEF3C7","#92400E","#F59E0B","🟡"),
        "red":    ("#FEE2E2","#991B1B","#EF4444","🔴"),
    }
    bg, fg, border, icon = color_cfg.get(r["color"], ("#F1F5F9","#334155","#94A3B8","📋"))

    st.progress(1.0, text="✅ 분류 완료!")
    st.write("")

    # ── st.metric 대시보드 상단 ────────────────────────────────────────────────
    m1, m2, m3 = st.columns(3)
    ecl_val = "✅ 적용" if r["ecl"] else "❌ 미적용"
    recycling_val = "✅ 발생" if r["recycling"] else ("⚠️ 금지" if r.get("recycling_note") else "❌ 없음")
    with m1:
        st.metric(label="최종 분류", value=r["classification"],
                  delta=r["label"].split("—")[0].strip() if "—" in r["label"] else "")
    with m2:
        st.metric(label="ECL 손상 인식", value=ecl_val)
    with m3:
        st.metric(label="OCI Recycling", value=recycling_val)

    st.write("")

    # ── 분류 결과 배너 ────────────────────────────────────────────────────────
    st.markdown(
        f'<div style="background:{bg};border:2px solid {border};border-radius:14px;'
        f'padding:1.3rem 1.7rem;margin-bottom:1.2rem">'
        f'<div style="font-size:1.5rem;font-weight:700;color:{fg};margin-bottom:.35rem">'
        f'{icon} {r["label"]}</div>'
        f'<div style="font-size:0.9rem;color:{fg};line-height:1.7;opacity:0.95">{r["reason"]}</div>'
        f'</div>',
        unsafe_allow_html=True,
    )

    # AI 모드 배지
    if st.session_state.get("ai_mode") and st.session_state.get("ai_confirmed"):
        st.info("🤖 **AI 하이브리드 분류**: STEP 0~2는 계약서 자동 분석, STEP 3(사업모형)은 AI 추론값이 자동 반영된 결과입니다.", icon="🤖")

    # ── SPPI 집중검토 노트 (v7 신규) ─────────────────────────────────────────
    sppi_note_ss = st.session_state.get("sppi_attention_note")
    if sppi_note_ss:
        st.error(sppi_note_ss, icon="🔍")

    # ── RCPS·지분상품 특별 안내 배너 ─────────────────────────────────────────
    if st.session_state.get("equity_rcps_hint"):
        eq_kws_display = ", ".join(st.session_state.get("equity_rcps_keywords", [])[:5])
        st.warning(
            f"⚠️ **지분성 키워드 감지**: 문서에서 `{eq_kws_display}` 등이 발견되었습니다.\n\n"
            "RCPS·전환우선주 등은 계약 조건에 따라 **채무상품(복합계약)** 또는 **지분상품**으로 분류될 수 있습니다.\n\n"
            "- 주된 의무가 원금+이자 상환이고 전환권이 부가된 경우 → **복합계약(채무)** 경로 적용\n"
            "- 지분 성격이 주된 경우(배당·잔여재산 우선권 중심) → **지분상품** 경로 적용, **FVOCI 취소불가 선택권** 검토\n\n"
            "분류가 불확실하다면 **수동 분류 모드**에서 'STEP A-2 FVOCI 취소불가 선택권' 단계를 직접 진행하세요.",
            icon="⚠️",
        )

    # ── BM 추론 점수 요약 (AI 모드일 때) ────────────────────────────────────
    bm_inf_ss = st.session_state.get("bm_inference")
    if st.session_state.get("ai_confirmed") and bm_inf_ss:
        conf_label_map = {"high": "🟢 높음", "medium": "🟡 중간", "low": "🔴 낮음"}
        bm_ko_map = {"hold": "AC", "both": "FVOCI", "trading": "FVPL", "ambiguous": "추가검토"}
        bm_display = bm_ko_map.get(bm_inf_ss.get("proposed_bm","?"), "?")
        conf_display = conf_label_map.get(bm_inf_ss.get("confidence","low"), "낮음")
        with st.expander(f"📊 사업모형 AI 추론 상세 — {bm_display} (신뢰도 {conf_display})", expanded=False):
            sc_c = st.columns(3)
            for col, (lbl, sc, clr) in zip(sc_c, [
                ("AC 신호", bm_inf_ss.get("ac_score",0), "#059669"),
                ("FVOCI 신호", bm_inf_ss.get("fvoci_score",0), "#3B82F6"),
                ("FVPL 신호", bm_inf_ss.get("fvpl_score",0), "#EF4444"),
            ]):
                with col:
                    total_s = max(bm_inf_ss.get("ac_score",0)+bm_inf_ss.get("fvoci_score",0)+bm_inf_ss.get("fvpl_score",0), 1)
                    pct_s = int(sc/total_s*100)
                    st.markdown(
                        f'<div style="text-align:center;font-size:0.75rem;color:#64748B;margin-bottom:3px">{lbl}</div>'
                        f'<div style="background:#E2E8F0;border-radius:4px;height:8px">'
                        f'<div style="background:{clr};width:{pct_s}%;height:8px;border-radius:4px"></div></div>'
                        f'<div style="text-align:center;font-size:0.72rem;color:{clr};font-weight:600;margin-top:2px">'
                        f'{sc:.1f}점 ({pct_s}%)</div>',
                        unsafe_allow_html=True,
                    )
            if bm_inf_ss.get("evidence_lines"):
                st.markdown("**감지된 근거 문구**")
                for ev in bm_inf_ss["evidence_lines"][:4]:
                    sig_c = {"AC":"#065F46","FVOCI":"#1E40AF","FVPL":"#991B1B"}.get(ev["signal"],"#475569")
                    st.markdown(
                        f'<div style="border-left:3px solid {sig_c};padding:4px 8px;margin-bottom:4px;'
                        f'background:#F8FAFC;border-radius:0 6px 6px 0;font-size:0.78rem">'
                        f'<b style="color:{sig_c}">{ev["signal"]}</b> · [{ev["role"]}] '
                        f'키워드: <b>{ev["keyword"]}</b><br>'
                        f'<span style="color:#64748B;font-style:italic">"{ev["text"]}"</span></div>',
                        unsafe_allow_html=True,
                    )

    st.divider()

    # ── 적용 기준서 조항 ──────────────────────────────────────────────────────
    st.markdown("### 📚 적용 기준서 조항")
    ref_cols = st.columns(min(len(r["refs"]), 5))
    for i, ref in enumerate(r["refs"]):
        ref_cols[i % 5].code(ref, language=None)

    st.divider()

    # ── 회계처리 요약표 (st.dataframe) ───────────────────────────────────────
    st.markdown("### 📋 회계처리 요약표")
    rows = _get_accounting_rows(r["classification"], r)
    df = pd.DataFrame(rows, columns=["항목", "내용"])
    st.dataframe(df, use_container_width=True, hide_index=True,
                 column_config={"항목": st.column_config.TextColumn(width="small"),
                                "내용": st.column_config.TextColumn(width="large")})

    st.divider()

    # ── ECL · Recycling ──────────────────────────────────────────────────────
    st.markdown("### ⚙️ 추가 처리 사항")
    if r["ecl"]:
        st.success("**ECL(기대신용손실) 손상 모형 적용 대상**\n\n매 보고기간 말에 기대신용손실을 산정하고 손실충당금을 인식해야 합니다. [§5.5]")
    else:
        st.info("**ECL(손상) 적용 없음** — 이 분류에는 손상 인식 규정이 적용되지 않습니다.")

    if r["recycling"]:
        st.warning("**채무상품 FVOCI**: 처분 시 OCI에 누적된 손익이 당기손익(P&L)으로 재분류됩니다. [§5.7.2]\n\n지분상품 FVOCI(Recycling 금지)와의 핵심 차이입니다.", icon="⚠️")
    elif r.get("recycling_note"):
        st.warning(r["recycling_note"], icon="⚠️")
    else:
        st.info("**Recycling 없음** — OCI 잔액은 처분 후에도 당기손익으로 이전되지 않습니다.")

    if r["warning"]:
        st.warning(f"**주의사항**\n\n{r['warning']}", icon="⚠️")

    st.divider()

    # ── SPPI 유사 사례 ────────────────────────────────────────────────────────
    case_key = r.get("case_key")
    if case_key and case_key in SPPI_CASES_DICT:
        case = SPPI_CASES_DICT[case_key]
        st.markdown("### 🗂️ 유사 사례 참고")
        c1, c2 = st.columns(2)
        with c1:
            st.markdown(f"**상품명**: {case['label']}")
            st.markdown(f"**유형**: `{case['category']}`")
            st.markdown(f"**기준서**: {', '.join(case['standard_ref'])}")
        with c2:
            st.markdown(f"**불충족 이유**: {case['reason']}")
            st.markdown(f"**판단 기준**: {case['judgment_criteria']}")
        with st.expander("📄 JSON 데이터 보기 (DB 확장용)", expanded=False):
            st.json(case)
        st.divider()

    # ── AI 최종 금융자산 분류 제안 섹션 (v7 신규) ──────────────────────────────
    if st.session_state.get("ai_confirmed"):
        _render_ai_final_conclusion(r, ans)

    # ── 입력 경로 요약 ────────────────────────────────────────────────────────
    st.markdown("### 🗺️ 분류 경로 요약")
    path_cols = st.columns(3)
    for i, (k, v) in enumerate(ans.items()):
        path_cols[i % 3].code(f"{_STEP_LABEL.get(k,k)}: {_VAL_KO.get(v,v)}", language=None)

    st.write("")
    c1, c2, c3 = st.columns(3)
    with c1:
        if st.button("← 이전 단계로", use_container_width=True):
            st.session_state.show_result = False
            seq = get_step_sequence(ans)
            if seq: ans.pop(seq[-1], None)
            if st.session_state.get("ai_confirmed"):
                ans.pop("s_fvo", None)
                ans.pop("s_bm", None)
            st.rerun()
    with c2:
        if st.button("🔄 처음부터 다시", type="primary", use_container_width=True):
            _full_reset()
    with c3:
        if st.button("📋 텍스트 리포트", use_container_width=True):
            _show_text_report(r, ans)


def _render_ai_final_conclusion(r: dict, ans: dict):
    """AI가 제안하는 최종 금융자산 분류 결론 섹션 (v7 신규)"""
    cls = r["classification"]
    label = r["label"]
    bm_inf_ss = st.session_state.get("bm_inference") or {}

    # 분류별 결론 문구 및 기준서 근거
    _CONCLUSIONS = {
        "AC": {
            "text": f"본 자산은 **상각후원가(AC) 측정 금융자산**으로 분류하는 것이 적절해 보입니다.",
            "basis": "계약상 현금흐름 수취 목적 사업모형(§B4.1.2C) + SPPI 충족(§4.1.2⑵) → AC 분류(§4.1.2)",
            "std": "§4.1.2",
            "ecl_note": "매 보고기간 말 ECL(기대신용손실)을 인식해야 합니다 [§5.5].",
            "recycling_note": None,
            "color": "#D1FAE5", "border": "#059669", "fg": "#065F46", "icon": "✅",
        },
        "FVOCI_DEBT": {
            "text": f"본 자산은 **기타포괄손익-공정가치(FVOCI) 측정 채무상품**으로 분류하는 것이 적절해 보입니다.",
            "basis": "수취+매도 병행 사업모형(§B4.1.4A) + SPPI 충족(§4.1.2A⑵) → FVOCI 분류(§4.1.2A)",
            "std": "§4.1.2A",
            "ecl_note": "ECL 적용 대상이며, 처분 시 OCI 누적손익이 P&L로 재분류(Recycling)됩니다 [§5.7.2].",
            "recycling_note": "처분 시 Recycling 발생 (채무상품 FVOCI)",
            "color": "#DBEAFE", "border": "#3B82F6", "fg": "#1E40AF", "icon": "🔵",
        },
        "FVOCI_EQ": {
            "text": f"본 자산은 **기타포괄손익-공정가치(FVOCI) 측정 지분상품**으로 분류하는 것이 적절해 보입니다.",
            "basis": "비단기매매 지분상품 FVOCI 취소불가 선택권 행사(§4.1.4) → 처분 시 Recycling 금지(§5.7.5)",
            "std": "§4.1.4 / §5.7.5",
            "ecl_note": "ECL 적용 없음. 처분 시에도 OCI 손익이 P&L로 이전되지 않습니다.",
            "recycling_note": "처분 시 Recycling 금지 (지분상품 FVOCI)",
            "color": "#DBEAFE", "border": "#3B82F6", "fg": "#1E40AF", "icon": "🔵",
        },
        "FVPL": {
            "text": f"본 자산은 **당기손익-공정가치(FVPL) 측정 금융자산**으로 분류하는 것이 적절해 보입니다.",
            "basis": "SPPI 불충족 또는 공정가치 실현 목적 사업모형 → FVPL 분류(§4.1.4 / §B4.1.5)",
            "std": "§4.1.4",
            "ecl_note": "ECL 적용 없음. 모든 공정가치 변동이 당기손익으로 인식됩니다.",
            "recycling_note": None,
            "color": "#FEE2E2", "border": "#EF4444", "fg": "#991B1B", "icon": "🔴",
        },
    }

    # 분류 키 결정
    if cls == "FVOCI":
        key = "FVOCI_EQ" if "지분" in label else "FVOCI_DEBT"
    elif cls == "FVPL":
        key = "FVPL"
    elif cls == "AC":
        key = "AC"
    else:
        key = "FVPL"

    c = _CONCLUSIONS[key]

    st.markdown("---")
    st.markdown(
        f'<div style="background:{NAVY_LIGHT};border:2px solid {NAVY};border-radius:14px;'
        f'padding:1.1rem 1.5rem;margin-bottom:1rem">'
        f'<div style="font-size:1rem;font-weight:700;color:{NAVY};margin-bottom:0.4rem">'
        f'🤖 AI가 제안하는 최종 회계처리</div>'
        f'<div style="font-size:0.85rem;color:#475569">문서 분석 및 분류 로직을 종합한 AI 최종 제안입니다. '
        f'전문가 검토 후 확정하시기 바랍니다.</div></div>',
        unsafe_allow_html=True,
    )

    # 결론 카드
    st.markdown(
        f'<div style="background:{c["color"]};border:2px solid {c["border"]};border-radius:12px;'
        f'padding:1rem 1.4rem;margin-bottom:0.75rem">'
        f'<div style="font-size:1.1rem;font-weight:700;color:{c["fg"]};margin-bottom:0.35rem">'
        f'{c["icon"]} {c["text"]}</div>'
        f'<div style="font-size:0.83rem;color:{c["fg"]};opacity:0.9;margin-bottom:0.4rem">'
        f'📌 근거: {c["basis"]}</div>'
        f'<div style="font-size:0.8rem;color:{c["fg"]};opacity:0.85">'
        f'⚙️ {c["ecl_note"]}</div>'
        + (f'<div style="font-size:0.8rem;color:{c["fg"]};opacity:0.85;margin-top:3px">'
           f'🔄 {c["recycling_note"]}</div>' if c.get("recycling_note") else "")
        + f'</div>',
        unsafe_allow_html=True,
    )

    # 기준서 조항 태그
    cols = st.columns(4)
    with cols[0]: st.code(c["std"], language=None)
    bm_used = ans.get("s_bm", "")
    bm_ko_m = {"hold": "AC 모형(§4.1.2)", "both": "FVOCI 모형(§4.1.2A)",
               "trading": "FVPL 잔여범주(§B4.1.5)", "ambiguous": "추가 검토 필요"}
    with cols[1]: st.code(bm_ko_m.get(bm_used, "—"), language=None)

    # 사업모형 점수 요약 (작게)
    if bm_inf_ss:
        total_s = max(bm_inf_ss.get("ac_score",0)+bm_inf_ss.get("fvoci_score",0)+bm_inf_ss.get("fvpl_score",0),1)
        winner_sig = {"hold":"AC","both":"FVOCI","trading":"FVPL"}.get(bm_used,"?")
        winner_score = {"hold":bm_inf_ss.get("ac_score",0),"both":bm_inf_ss.get("fvoci_score",0),
                        "trading":bm_inf_ss.get("fvpl_score",0)}.get(bm_used,0)
        pct = int(winner_score/total_s*100)
        with cols[2]:
            st.markdown(
                f'<div style="text-align:center;font-size:0.72rem;color:#64748B">BM 점수</div>'
                f'<div style="font-size:0.85rem;font-weight:600;color:{NAVY};text-align:center">'
                f'{winner_sig} {winner_score:.1f}점 ({pct}%)</div>',
                unsafe_allow_html=True,
            )
        conf_disp = {"high":"🟢 높음","medium":"🟡 보통","low":"🔴 낮음"}.get(
            bm_inf_ss.get("confidence","low"),"?")
        with cols[3]:
            st.markdown(
                f'<div style="text-align:center;font-size:0.72rem;color:#64748B">신뢰도</div>'
                f'<div style="font-size:0.85rem;font-weight:600;text-align:center">{conf_disp}</div>',
                unsafe_allow_html=True,
            )


def _get_accounting_rows(cls: str, r: dict) -> list:
    is_eq_fvoci = cls == "FVOCI" and "지분" in r["label"]
    if cls == "AC":
        return [
            ("최초 인식","공정가치 + 거래원가"),
            ("후속 측정","유효이자율법 상각후원가"),
            ("이자수익 반영","당기손익 (P&L) — 유효이자율법"),
            ("평가손익 반영","해당 없음 (공정가치 변동 미반영)"),
            ("ECL 손상 인식","✅ 적용 — 12개월 또는 전체기간 ECL [§5.5]"),
            ("처분 손익","당기손익 (P&L)"),
            ("OCI Recycling","해당 없음"),
        ]
    if cls == "FVOCI" and not is_eq_fvoci:
        return [
            ("최초 인식","공정가치"),
            ("후속 측정","공정가치"),
            ("이자수익 반영","당기손익 (P&L) — 유효이자율법"),
            ("평가손익 반영","기타포괄손익 (OCI)"),
            ("ECL 손상 인식","✅ 적용 — P&L에 반영 (OCI 조정 포함) [§5.5]"),
            ("처분 시 OCI 재분류","✅ P&L로 Recycling 발생 [§5.7.2]"),
            ("OCI Recycling","✅ 처분 시 발생"),
        ]
    if cls == "FVOCI" and is_eq_fvoci:
        return [
            ("최초 인식","공정가치"),
            ("후속 측정","공정가치"),
            ("배당 반영","당기손익 (P&L) [§B5.7.1]"),
            ("평가손익 반영","기타포괄손익 (OCI)"),
            ("ECL 손상 인식","❌ 미적용"),
            ("처분 시 OCI 재분류","❌ P&L 재분류 금지 — 자본 내 이전만 허용 [§5.7.5]"),
            ("OCI Recycling","❌ 금지 (지분상품 FVOCI)"),
        ]
    return [
        ("최초 인식","공정가치"),
        ("후속 측정","공정가치"),
        ("이자/배당 반영","당기손익 (P&L)"),
        ("평가손익 반영","당기손익 (P&L) — 전액"),
        ("ECL 손상 인식","❌ 미적용"),
        ("처분 손익","당기손익 (P&L)"),
        ("OCI Recycling","해당 없음"),
    ]


def _show_text_report(r: dict, ans: dict):
    lines = [
        "="*65, "  IFRS 9 금융자산 분류 결과 리포트", "="*65,
        f"  분류: {r['label']}", "",
        "[ 분류 근거 ]", f"  {r['reason']}", "",
        "[ 적용 기준서 조항 ]", "  " + "  |  ".join(r["refs"]), "",
        "[ 후속 측정 원칙 ]",
    ]
    for item in r["accounting"]: lines.append(f"  ✦ {item}")
    lines += ["",
              f"[ ECL 손상 적용 ]   {'✅ 적용 (§5.5)' if r['ecl'] else '❌ 미적용'}",
              f"[ Recycling 발생 ]  {'✅ 처분 시 발생 (채무 FVOCI)' if r['recycling'] else '❌ 없음'}"]
    if r.get("recycling_note"): lines.append(f"  ※ {r['recycling_note']}")
    if r["warning"]: lines += ["", "[ 주의사항 ]", f"  ⚠️  {r['warning']}"]
    lines += ["", "[ 분류 경로 ]"]
    for k, v in ans.items(): lines.append(f"  {_STEP_LABEL.get(k,k)}: {_VAL_KO.get(v,v)}")
    if st.session_state.get("ai_mode") and st.session_state.get("ai_confirmed"):
        lines += ["", "[ 분류 방식 ]", "  🤖 AI 하이브리드 (STEP 0~2 자동 + STEP 3 수동)"]
    lines += ["", "="*65]
    st.code("\n".join(lines), language=None)


# ══════════════════════════════════════════════════════════════════════════════
# 14. 헬퍼 함수
# ══════════════════════════════════════════════════════════════════════════════

def _pick(step_id: str, val: str):
    st.session_state.answers[step_id] = val
    st.rerun()


def _go_next():
    ans = st.session_state.answers
    seq = get_step_sequence(ans)
    if ans.get(seq[-1]):
        st.session_state.history.append(seq[-1])
        if is_terminal(ans):
            st.session_state.show_result = True
        st.rerun()


def _go_back():
    if not st.session_state.history: return
    prev_id = st.session_state.history.pop()
    seq = get_step_sequence(st.session_state.answers)
    if prev_id in seq:
        for sid in seq[seq.index(prev_id):]:
            st.session_state.answers.pop(sid, None)
    st.session_state.show_result = False
    st.rerun()


# ══════════════════════════════════════════════════════════════════════════════
# 15. main()
# ══════════════════════════════════════════════════════════════════════════════

def main():
    st.set_page_config(
        page_title="IFRS 9 금융자산 분류 마법사",
        page_icon="📊",
        layout="wide",
        initial_sidebar_state="expanded",
    )

    _init_session()  # 세션 상태 전수 초기화 (누락 방지)
    _inject_css()    # 전역 CSS 주입

    # ── 사이드바 렌더링 ───────────────────────────────────────────────────────
    _render_sidebar()

    # ── 메인 헤더 ─────────────────────────────────────────────────────────────
    st.markdown(
        f'<h1 style="color:{NAVY};font-weight:700;letter-spacing:-0.03em;margin-bottom:0.2rem">'
        f'📊 IFRS 9 금융자산 분류 마법사</h1>',
        unsafe_allow_html=True,
    )
    st.caption("K-IFRS 1109호 4장 · 부록B · PwC 실무 가이드라인 | FVPL · AC · FVOCI | 모든 결과에 **§ 조항** 포함")
    st.divider()

    ans = st.session_state.answers

    # ══════════════════════════════════════════════════════════════════════════
    # 화면 분기
    # ══════════════════════════════════════════════════════════════════════════

    # Case 1: 시작 전 환영 화면
    if not st.session_state.wizard_started:
        _render_welcome()
        return

    # Case 2: 결과 화면
    if st.session_state.show_result:
        _render_result(ans)
        return

    # Case 3: AI 모드 — 분석 결과 확인 화면
    if st.session_state.ai_mode and not st.session_state.ai_confirmed:
        if st.session_state.ai_result:
            _render_ai_confirm()
        else:
            st.error("AI 분석 결과가 없습니다. 사이드바에서 파일을 다시 업로드해주세요.")
        return

    # Case 4: AI 모드 — 사업모형 단계 (직접 선택)
    if st.session_state.ai_mode and st.session_state.ai_confirmed:
        if "s_bm" not in ans:
            _render_progress(ans)
            st.write("")
            _render_bm_after_ai()
            return
        # s_bm 답 있고 hold/both → FVO 단계
        bm = ans.get("s_bm")
        if bm in ("hold", "both") and "s_fvo" not in ans:
            _render_progress(ans)
            st.write("")
            _render_step("s_fvo", ans)
            return
        # 모든 답 완료 → 결과
        st.session_state.show_result = True
        st.rerun()
        return

    # Case 5: 수동 마법사 모드
    seq = get_step_sequence(ans)
    current_step_id = seq[-1]

    if is_terminal(ans):
        st.session_state.show_result = True
        st.rerun()
        return

    _render_progress(ans)
    st.write("")
    _render_step(current_step_id, ans)


# ══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    main()
